"""Multi-format export adapters for MEAL registers and reports.

CSV adapters mirror ``apps.programs.exports``; XLSX/DOCX/PDF adapters mirror
``apps.programs.report_exports``.  Every export is permission checked and
produces a structured audit log entry.
"""

from __future__ import annotations

import csv
import logging
from io import BytesIO, StringIO
from pathlib import Path

from django.conf import settings
from django.core.exceptions import PermissionDenied
from django.http import HttpResponse

from apps.rbac.authorization import user_has_permission

from .models import Evaluation, Indicator, LessonLearned, MEALReport, MonitoringVisit
from .permissions import MEAL_EXPORT, MEAL_MANAGE
from .selectors import (
    meal_queryset,
    user_can_access_meal_record,
    visible_complaints,
    visible_feedback,
)

logger = logging.getLogger(__name__)

LOGO_PATH = Path(settings.BASE_DIR) / "static" / "images" / "app_logo.png"


def formula_safe_csv_value(value) -> str:
    text = str(value or "")
    if text.startswith(("=", "+", "-", "@", "\t", "\r")):
        return f"'{text}"
    return text


def _require_export_permission(user) -> None:
    if not (
        user_has_permission(user, MEAL_EXPORT) or user_has_permission(user, MEAL_MANAGE)
    ):
        raise PermissionDenied("MEAL export permission is required.")


def _csv_response(user, headers, rows, filename: str, event: str) -> HttpResponse:
    output = StringIO(newline="")
    writer = csv.writer(output)
    writer.writerow(headers)
    for row in rows:
        writer.writerow([formula_safe_csv_value(value) for value in row])
    response = HttpResponse(output.getvalue(), content_type="text/csv; charset=utf-8")
    response["Content-Disposition"] = f'attachment; filename="{filename}"'
    response["Cache-Control"] = "private, no-store"
    response["Pragma"] = "no-cache"
    response["X-Content-Type-Options"] = "nosniff"
    logger.info(
        event,
        extra={
            "meal_event": {
                "action": event,
                "format": "csv",
                "actor_id": str(user.pk),
                "row_count": len(rows),
                "scope": "visible_meal_records",
            }
        },
    )
    return response


def _visible_indicators(user):
    return meal_queryset(user, Indicator)


def indicator_register_csv_response(user) -> HttpResponse:
    _require_export_permission(user)
    headers = [
        "Indicator ID",
        "Code",
        "Title",
        "Type",
        "Category",
        "Unit",
        "Status",
        "Responsible officer",
    ]
    rows = []
    for indicator in _visible_indicators(user).select_related(
        "category", "unit", "responsible_officer"
    ):
        rows.append(
            [
                indicator.reference_number,
                indicator.code,
                indicator.title,
                indicator.get_indicator_type_display(),
                indicator.category.name if indicator.category else "",
                indicator.unit.name if indicator.unit else "",
                indicator.get_status_display(),
                (
                    indicator.responsible_officer.full_name
                    if indicator.responsible_officer
                    else ""
                ),
            ]
        )
    return _csv_response(
        user, headers, rows, "indicator_registry.csv", "indicator_register_exported"
    )


def monitoring_visit_register_csv_response(user) -> HttpResponse:
    _require_export_permission(user)
    headers = [
        "Visit ID",
        "Title",
        "Program",
        "Project",
        "Visit date",
        "Community",
        "Status",
        "Follow-up due",
    ]
    rows = []
    for visit in meal_queryset(user, MonitoringVisit).select_related(
        "program", "project"
    ):
        rows.append(
            [
                visit.reference_number,
                visit.reference_number,
                visit.program.reference_number if visit.program else "",
                visit.project.reference_number if visit.project else "",
                visit.visit_date,
                visit.community,
                visit.get_status_display(),
                visit.follow_up_due or "",
            ]
        )
    return _csv_response(
        user,
        headers,
        rows,
        "monitoring_visit_register.csv",
        "monitoring_visit_exported",
    )


def complaint_register_csv_response(user) -> HttpResponse:
    _require_export_permission(user)
    headers = [
        "Complaint ID",
        "Submission date",
        "Category",
        "Channel",
        "Priority",
        "Status",
        "Confidential",
    ]
    rows = []
    for complaint in visible_complaints(user).select_related("category", "channel"):
        rows.append(
            [
                complaint.reference_number,
                complaint.submission_date,
                complaint.category.name if complaint.category else "",
                complaint.channel.name if complaint.channel else "",
                complaint.get_priority_display(),
                complaint.get_status_display(),
                "Yes" if complaint.is_confidential else "No",
            ]
        )
    return _csv_response(
        user, headers, rows, "complaint_register.csv", "complaint_register_exported"
    )


def feedback_register_csv_response(user) -> HttpResponse:
    _require_export_permission(user)
    headers = [
        "Feedback ID",
        "Submission date",
        "Category",
        "Channel",
        "Status",
        "Rating",
        "Confidential",
    ]
    rows = []
    for feedback in visible_feedback(user).select_related("category", "channel"):
        rows.append(
            [
                feedback.reference_number,
                feedback.submission_date,
                feedback.category.name if feedback.category else "",
                feedback.channel.name if feedback.channel else "",
                feedback.get_status_display(),
                feedback.satisfaction_rating or "",
                "Yes" if feedback.is_confidential else "No",
            ]
        )
    return _csv_response(
        user, headers, rows, "feedback_register.csv", "feedback_register_exported"
    )


def evaluation_register_csv_response(user) -> HttpResponse:
    _require_export_permission(user)
    headers = [
        "Evaluation ID",
        "Title",
        "Program",
        "Project",
        "Start date",
        "End date",
        "Status",
    ]
    rows = []
    for evaluation in meal_queryset(user, Evaluation).select_related(
        "program", "project"
    ):
        rows.append(
            [
                evaluation.reference_number,
                evaluation.title,
                evaluation.program.reference_number if evaluation.program else "",
                evaluation.project.reference_number if evaluation.project else "",
                evaluation.start_date,
                evaluation.end_date or "",
                evaluation.get_status_display(),
            ]
        )
    return _csv_response(
        user, headers, rows, "evaluation_register.csv", "evaluation_register_exported"
    )


def meal_report_register_csv_response(user) -> HttpResponse:
    _require_export_permission(user)
    headers = [
        "Report ID",
        "Title",
        "Type",
        "Period start",
        "Period end",
        "Status",
        "Prepared by",
    ]
    rows = []
    for report in meal_queryset(user, MEALReport).select_related("prepared_by"):
        rows.append(
            [
                report.reference_number,
                report.title,
                report.get_report_type_display(),
                report.period_start or "",
                report.period_end or "",
                report.get_status_display(),
                report.prepared_by.full_name if report.prepared_by else "",
            ]
        )
    return _csv_response(
        user, headers, rows, "meal_report_register.csv", "meal_report_register_exported"
    )


# ── XLSX / DOCX / PDF report export ───────────────────────────────────────


def _xlsx_styled_header(ws, headers: list) -> None:
    from openpyxl.styles import (  # type: ignore[import-untyped]
        Alignment,
        Border,
        Font,
        PatternFill,
        Side,
    )

    fill = PatternFill(start_color="003366", end_color="003366", fill_type="solid")
    font = Font(color="FFFFFF", bold=True, size=11)
    border = Border(
        left=Side(style="thin"),
        right=Side(style="thin"),
        top=Side(style="thin"),
        bottom=Side(style="thin"),
    )
    for col_idx, header in enumerate(headers, 1):
        cell = ws.cell(row=1, column=col_idx, value=header)
        cell.font = font
        cell.fill = fill
        cell.border = border
        cell.alignment = Alignment(horizontal="center", wrap_text=True)


def _xlsx_set_widths(ws, widths: list) -> None:
    from openpyxl.utils import get_column_letter  # type: ignore[import-untyped]

    for col_idx, width in enumerate(widths, 1):
        ws.column_dimensions[get_column_letter(col_idx)].width = width


def _build_xlsx(rows: list[list], headers: list, widths: list) -> BytesIO:
    from openpyxl import Workbook  # type: ignore[import-untyped]

    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "MEAL Report"
    _xlsx_styled_header(sheet, headers)
    for row in rows:
        sheet.append(row)
    _xlsx_set_widths(sheet, widths)
    buffer = BytesIO()
    workbook.save(buffer)
    buffer.seek(0)
    return buffer


def _build_docx(title: str, rows: list[list], headers: list) -> BytesIO:
    from docx import Document

    document = Document()
    document.add_heading(title, level=0)
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Light Grid Accent 1"
    for col_idx, header in enumerate(headers):
        table.rows[0].cells[col_idx].text = str(header)
    for row in rows:
        cells = table.add_row().cells
        for col_idx, value in enumerate(row):
            cells[col_idx].text = str(value)
    buffer = BytesIO()
    document.save(buffer)
    buffer.seek(0)
    return buffer


def _build_pdf(title: str, rows: list[list], headers: list) -> BytesIO:
    from reportlab.lib import colors  # type: ignore[import-untyped]
    from reportlab.lib.pagesizes import A4, landscape  # type: ignore[import-untyped]
    from reportlab.lib.styles import getSampleStyleSheet  # type: ignore[import-untyped]
    from reportlab.platypus import (  # type: ignore[import-untyped]
        SimpleDocTemplate,
        Table,
        TableStyle,
    )

    buffer = BytesIO()
    doc = SimpleDocTemplate(
        buffer,
        pagesize=landscape(A4),
        title=title,
    )
    styles = getSampleStyleSheet()
    elements = [_reportlab_heading(title, styles)]
    data = [headers] + [[str(value) for value in row] for row in rows]
    table = Table(data, repeatRows=1)
    table.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#003366")),
                ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                ("FONTSIZE", (0, 0), (-1, 0), 8),
                ("FONTSIZE", (0, 1), (-1, -1), 7),
                ("GRID", (0, 0), (-1, -1), 0.25, colors.grey),
                (
                    "ROWBACKGROUNDS",
                    (0, 1),
                    (-1, -1),
                    [colors.white, colors.HexColor("#EEF2F7")],
                ),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
            ]
        )
    )
    elements.append(table)
    doc.build(elements)
    buffer.seek(0)
    return buffer


def _reportlab_heading(title: str, styles) -> object:
    from reportlab.lib import colors
    from reportlab.lib.styles import ParagraphStyle
    from reportlab.platypus import Paragraph

    style = ParagraphStyle(
        "MealTitle",
        parent=styles["Title"],
        fontName="Helvetica-Bold",
        fontSize=14,
        leading=18,
        textColor=colors.HexColor("#003366"),
    )
    return Paragraph(title, style)


def meal_report_export_response(user, report, fmt: str) -> HttpResponse:
    """Export a single MEAL report in the requested format."""
    if not user_can_access_meal_record(user, report):
        raise PermissionDenied("You do not have access to this report.")
    _require_export_permission(user)

    title = f"{report.reference_number} - {report.title}"
    headers = ["Field", "Value"]
    rows = [
        ["Report ID", report.reference_number],
        ["Title", report.title],
        ["Report type", report.get_report_type_display()],
        ["Program", report.program.title if report.program else ""],
        ["Project", report.project.title if report.project else ""],
        ["Period", f"{report.period_start} to {report.period_end}"],
        ["Status", report.get_status_display()],
        ["Prepared by", report.prepared_by.full_name if report.prepared_by else ""],
        ["Approved at", report.approved_at or ""],
        ["Content", report.content],
    ]
    widths = [28, 120]

    if fmt == "xlsx":
        buffer = _build_xlsx(rows, headers, widths)
        content_type = (
            "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
        )
        filename = f"{report.reference_number}.xlsx"
    elif fmt == "docx":
        buffer = _build_docx(title, rows, headers)
        content_type = (
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
        filename = f"{report.reference_number}.docx"
    elif fmt == "pdf":
        buffer = _build_pdf(title, rows, headers)
        content_type = "application/pdf"
        filename = f"{report.reference_number}.pdf"
    else:
        raise PermissionDenied(f"Unsupported export format: {fmt}")

    response = HttpResponse(buffer.getvalue(), content_type=content_type)
    response["Content-Disposition"] = f'attachment; filename="{filename}"'
    response["Cache-Control"] = "private, no-store"
    response["Pragma"] = "no-cache"
    response["X-Content-Type-Options"] = "nosniff"
    logger.info(
        "meal_report_exported",
        extra={
            "meal_event": {
                "action": "meal.report_exported",
                "format": fmt,
                "actor_id": str(user.pk),
                "entity_type": "MEALReport",
                "entity_id": str(report.pk),
            }
        },
    )
    return response


def lesson_register_csv_response(user) -> HttpResponse:
    _require_export_permission(user)
    headers = [
        "Lesson ID",
        "Title",
        "Category",
        "Status",
        "Responsible team",
    ]
    rows = []
    for lesson in meal_queryset(user, LessonLearned).select_related("category"):
        rows.append(
            [
                lesson.reference_number,
                lesson.title,
                lesson.category.name if lesson.category else "",
                lesson.get_status_display(),
                lesson.responsible_team,
            ]
        )
    return _csv_response(
        user, headers, rows, "lessons_learned_register.csv", "lesson_register_exported"
    )
