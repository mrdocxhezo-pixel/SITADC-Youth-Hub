"""Export services for report instances.

Supports PDF, DOCX, XLSX, CSV, HTML, and JSON export formats.
Each format renders the actual Report Instance data (field responses
saved by the user) through the Template Schema structure, with the
official SITADC app logo centered at the top of every report.
"""

from __future__ import annotations

import csv
import io
import json
import logging
import os
import re
from typing import Any

from django.conf import settings
from django.http import HttpResponse
from django.utils import timezone
from django.utils.html import strip_tags

logger = logging.getLogger(__name__)


# ---------------------------------------------------------------------------
# Logo Resolution
# ---------------------------------------------------------------------------


def _resolve_logo_path() -> str | None:
    """Return the absolute filesystem path to static/images/app_logo.png."""
    staticfiles_dirs = getattr(settings, "STATICFILES_DIRS", None)
    candidates = []
    if staticfiles_dirs:
        for d in staticfiles_dirs:
            candidates.append(os.path.join(str(d), "images", "app_logo.png"))
    candidates.append(
        os.path.join(str(settings.BASE_DIR), "static", "images", "app_logo.png")
    )
    for candidate in candidates:
        if os.path.exists(candidate):
            return candidate
    return None


# ---------------------------------------------------------------------------
# Value Rendering Helpers
# ---------------------------------------------------------------------------


def _render_display_value(field, value, options_map=None):
    """Render a field value into a human-readable display string.

    Handles all field types: text, numbers, dates, choices, multi-select,
    booleans, tables, rich text, file references, etc.
    """
    if value is None or value == "":
        return ""

    field_type = getattr(field, "field_type", "")

    # Boolean fields
    if field_type in ("TOGGLE", "CHECKBOX"):
        if isinstance(value, bool):
            return "Yes" if value else "No"
        if isinstance(value, str):
            lower = value.lower().strip()
            if lower in ("true", "1", "yes"):
                return "Yes"
            elif lower in ("false", "0", "no", ""):
                return "No"
        return "Yes" if value else "No"

    # Multi-select / checkbox groups
    if field_type in ("MULTI_SELECT", "CHECKBOX") and isinstance(value, list):
        if options_map:
            labels = [options_map.get(v, str(v)) for v in value]
        else:
            labels = [str(v) for v in value]
        return ", ".join(labels) if labels else ""

    # Single selection (dropdown / radio)
    if field_type in ("DROPDOWN", "RADIO") and options_map:
        if value in options_map:
            return options_map[value]
        # Try case-insensitive match
        lower_val = str(value).lower()
        for k, v in options_map.items():
            if str(k).lower() == lower_val:
                return v

    # Table / grid fields
    if field_type == "TABLE_GRID" and isinstance(value, list):
        return value  # Will be rendered as a table separately

    # Rich text — strip HTML tags for plain exports, keep for rich exports
    if field_type == "RICH_TEXT" and isinstance(value, str):
        return strip_tags(value).strip()

    # File uploads — return filename
    if field_type in ("IMAGE", "VIDEO", "AUDIO", "DOCUMENT") and isinstance(value, str):
        # If it's a storage path, extract just the filename
        if "/" in value:
            return value.split("/")[-1]
        return value

    # Numeric fields — format nicely
    if field_type in ("INTEGER", "CURRENCY", "PERCENTAGE") and value is not None:
        try:
            num = int(float(value))
            if field_type == "CURRENCY":
                return f"K {num:,.2f}" if num >= 0 else f"-K {abs(num):,.2f}"
            if field_type == "PERCENTAGE":
                return f"{num}%"
            return f"{num:,}"
        except (ValueError, TypeError):
            return str(value)

    if field_type == "DECIMAL" and value is not None:
        try:
            num = float(value)
            return f"{num:,.2f}"
        except (ValueError, TypeError):
            return str(value)

    # Date fields
    if field_type == "DATE" and value:
        return str(value)
    if field_type == "TIME" and value:
        return str(value)
    if field_type == "DATETIME" and value:
        return str(value)

    # GPS coordinates
    if field_type == "GPS_COORDINATES" and isinstance(value, dict):
        lat = value.get("latitude", "")
        lon = value.get("longitude", "")
        if lat and lon:
            return f"{lat}, {lon}"

    # Default — convert to string
    return str(value)


def _get_options_map(field):
    """Build a {value: label} map from a DynamicField's options."""
    options = {}
    for opt in field.options.all():
        options[opt.value] = opt.label
    return options


# ---------------------------------------------------------------------------
# Structured Data Normalization
# ---------------------------------------------------------------------------


def _build_structured_data(report):
    """Build a normalized, schema-aware data structure for the report.

    Returns a dict with:
        - meta: report metadata (title, reference, status, etc.)
        - sections: ordered list of section dicts, each containing
          ordered field dicts with label, type, raw_value, display_value
        - table_responses: table field responses keyed by field pk
        - group_responses: repeating group responses
        - evidence: evidence items
        - attachments: attachment items
    """
    from apps.reports.models import DynamicField, FieldGroup, TemplateSection

    # Build field response lookup: field_id -> ReportFieldResponse
    field_response_map = {}
    for fr in report.field_responses.select_related("field").all():
        field_response_map[str(fr.field_id)] = fr

    # Build table response lookup: field_id -> ReportTableResponse
    table_response_map = {}
    for tr in report.table_responses.select_related("table_field").all():
        table_response_map[str(tr.table_field_id)] = tr

    # Build section response lookup: section_id -> ReportSectionResponse
    section_response_map = {}
    for sr in report.section_responses.select_related("section").all():
        section_response_map[str(sr.section_id)] = sr

    # Build group response lookup: group_id -> [ReportGroupResponse]
    group_response_map = {}
    for gr in report.group_responses.select_related("group").all():
        gid = str(gr.group_id)
        if gid not in group_response_map:
            group_response_map[gid] = []
        group_response_map[gid].append(gr)

    # Iterate through template schema structure
    sections = []
    template_sections = TemplateSection.objects.filter(
        template=report.template
    ).order_by("sort_order", "name")

    for section in template_sections:
        section_fields = []

        for group in section.groups.order_by("sort_order", "name"):
            for field in group.fields.order_by("sort_order", "label"):
                if getattr(field, "hidden", False):
                    continue

                options_map = _get_options_map(field)
                fr = field_response_map.get(str(field.pk))
                raw_value = fr.value if fr else None

                # For table fields, get from table_response
                if field.field_type == "TABLE_GRID":
                    tr = table_response_map.get(str(field.pk))
                    raw_value = tr.rows if tr else []

                display_value = _render_display_value(field, raw_value, options_map)

                section_fields.append({
                    "field_id": str(field.pk),
                    "label": field.label,
                    "field_type": field.field_type,
                    "required": field.required,
                    "help_text": field.help_text or "",
                    "raw_value": raw_value,
                    "display_value": display_value,
                    "options_map": options_map,
                    "sort_order": field.sort_order,
                })

        # Also include any repeating group data
        for group in section.groups.order_by("sort_order", "name"):
            group_instances = group_response_map.get(str(group.pk), [])
            if group_instances:
                # Collect all unique field keys from group data
                all_field_keys = set()
                for gi in group_instances:
                    if isinstance(gi.data, dict):
                        all_field_keys.update(gi.data.keys())

                if all_field_keys:
                    section_fields.append({
                        "field_id": f"__group__{group.pk}",
                        "label": group.name,
                        "field_type": "REPEATING_GROUP",
                        "required": False,
                        "help_text": group.description or "",
                        "raw_value": [gi.data for gi in sorted(group_instances, key=lambda x: x.instance_index)],
                        "display_value": "",
                        "options_map": {},
                        "sort_order": group.sort_order + 0.5,
                    })

        section_response = section_response_map.get(str(section.pk))

        sections.append({
            "section_id": str(section.pk),
            "name": section.name,
            "description": section.description or "",
            "fields": section_fields,
            "is_complete": section_response.is_complete if section_response else False,
        })

    # Evidence & attachments
    evidence = []
    for e in report.evidence_items.all():
        evidence.append({
            "type": e.get_evidence_type_display(),
            "filename": e.original_filename,
            "description": e.description,
            "verified": e.is_verified,
        })

    attachments = []
    for a in report.attachments.all():
        attachments.append({
            "filename": a.original_filename,
            "size": a.file_size,
            "description": a.description,
        })

    # Reporting period display
    reporting_period = ""
    if report.reporting_period:
        reporting_period = str(report.reporting_period)
    elif hasattr(report, "_reporting_period_display"):
        reporting_period = report._reporting_period_display

    meta = {
        "title": report.title,
        "reference_number": report.reference_number,
        "template": report.template.title if report.template else "",
        "template_code": report.template.code if report.template else "",
        "category": report.category.name if report.category else "",
        "status": report.get_status_display(),
        "validation_status": report.get_validation_status_display(),
        "confidentiality": report.get_confidentiality_display(),
        "owner": str(report.owner) if report.owner else "",
        "department": report.department or "",
        "due_date": str(report.due_date) if report.due_date else "",
        "submitted_at": str(report.submitted_at) if report.submitted_at else "",
        "approved_at": str(report.approved_at) if report.approved_at else "",
        "version_number": report.version_number,
        "notes": report.notes or "",
        "reporting_period": reporting_period,
        "is_draft": report.is_draft,
        "created_at": str(report.created_at),
        "updated_at": str(report.updated_at),
        "template_version": (
            report.template_version.version_number if report.template_version else ""
        ),
    }

    return {
        "meta": meta,
        "sections": sections,
        "evidence": evidence,
        "attachments": attachments,
    }


# ---------------------------------------------------------------------------
# HTML Export
# ---------------------------------------------------------------------------


def export_html(report: Any) -> HttpResponse:
    """Export report as HTML with centered logo and structured sections."""
    data = _build_structured_data(report)
    meta = data["meta"]

    logo_path = _resolve_logo_path()
    logo_html = ""
    if logo_path:
        logo_url = f"/static/images/app_logo.png"
        logo_html = (
            f'<div style="text-align:center; margin-bottom:20px;">'
            f'<img src="{logo_url}" alt="SITADC Logo" '
            f'style="max-width:180px; height:auto;">'
            f"</div>"
        )

    # Status badge
    status_color = {
        "Draft": "#6c757d",
        "In Progress": "#17a2b8",
        "Ready for Submission": "#0d6efd",
        "Submitted": "#0d6efd",
        "Under Review": "#fd7e14",
        "Pending Approval": "#ffc107",
        "Approved": "#198754",
        "Finalized": "#198754",
        "Archived": "#343a40",
        "Returned for Correction": "#dc3545",
        "Rejected": "#dc3545",
        "Validation Failed": "#dc3545",
    }
    badge_color = status_color.get(meta["status"], "#6c757d")

    # Draft banner
    draft_banner = ""
    if meta["is_draft"]:
        draft_banner = (
            '<div style="background:#B22222; color:white; text-align:center; '
            'padding:8px; font-weight:bold; font-size:12px; margin-bottom:20px;">'
            "DRAFT &mdash; NOT FOR OFFICIAL USE"
            "</div>"
        )

    html_parts = [
        "<!DOCTYPE html>",
        "<html><head><meta charset='utf-8'>",
        f"<title>{meta['title']}</title>",
        "<style>",
        "body { font-family: 'Segoe UI', Arial, sans-serif; margin: 40px; color: #333; }",
        ".report-header { text-align: center; margin-bottom: 30px; }",
        ".report-header h1 { color: #003366; font-size: 22px; margin: 15px 0 5px 0; }",
        ".report-header .subtitle { color: #555; font-size: 13px; }",
        ".meta-table { width: 100%; border-collapse: collapse; margin: 15px 0 25px 0; }",
        ".meta-table td { padding: 6px 10px; font-size: 13px; border-bottom: 1px solid #eee; }",
        ".meta-table td:first-child { font-weight: bold; color: #003366; width: 200px; background: #f8f9fa; }",
        "hr.section-divider { border: none; border-top: 2px solid #003366; margin: 30px 0 15px 0; }",
        "h2 { color: #003366; font-size: 16px; border-bottom: 1px solid #ddd; padding-bottom: 6px; margin-top: 25px; }",
        "h3 { color: #333; font-size: 14px; margin-top: 15px; }",
        ".field-row { margin: 8px 0; }",
        ".field-label { font-weight: bold; color: #444; font-size: 13px; }",
        ".field-value { color: #333; font-size: 13px; margin-top: 2px; }",
        ".empty-value { color: #999; font-style: italic; }",
        "table.data-table { border-collapse: collapse; width: 100%; margin: 10px 0; }",
        "table.data-table th, table.data-table td { border: 1px solid #ddd; padding: 6px 8px; text-align: left; font-size: 12px; }",
        "table.data-table th { background: #003366; color: white; font-weight: bold; }",
        "table.data-table tr:nth-child(even) { background: #f8f9fa; }",
        ".badge { padding: 3px 10px; border-radius: 4px; font-size: 12px; color: white; }",
        ".footer { margin-top: 40px; padding-top: 15px; border-top: 1px solid #ccc; font-size: 11px; color: #666; text-align: center; }",
        "</style></head><body>",
        # Logo
        logo_html,
        draft_banner,
        # Header
        '<div class="report-header">',
        f"<h1>{meta['title'].upper()}</h1>",
        "</div>",
        # Metadata
        '<table class="meta-table">',
        f'<tr><td>Category</td><td>{meta["category"]}</td></tr>',
        f'<tr><td>Template</td><td>{meta["template"]}</td></tr>',
        f'<tr><td>Reference</td><td>{meta["reference_number"]}</td></tr>',
    ]

    if meta["reporting_period"]:
        html_parts.append(
            f'<tr><td>Reporting Period</td><td>{meta["reporting_period"]}</td></tr>'
        )
    if meta["template_version"]:
        html_parts.append(
            f'<tr><td>Template Version</td><td>{meta["template_version"]}</td></tr>'
        )
    if meta["department"]:
        html_parts.append(
            f'<tr><td>Department</td><td>{meta["department"]}</td></tr>'
        )
    html_parts.append(
        f'<tr><td>Status</td><td><span class="badge" style="background:{badge_color};">{meta["status"]}</span></td></tr>'
    )
    if meta["owner"]:
        html_parts.append(f'<tr><td>Prepared By</td><td>{meta["owner"]}</td></tr>')
    html_parts.append(
        f'<tr><td>Generated</td><td>{timezone.now().strftime("%Y-%m-%d %H:%M")}</td></tr>'
    )
    html_parts.append("</table>")

    # Sections
    for section in data["sections"]:
        if not section["fields"]:
            continue

        html_parts.append(f'<hr class="section-divider">')
        html_parts.append(f"<h2>{section['name'].upper()}</h2>")

        if section["description"]:
            html_parts.append(
                f'<p style="color:#666; font-size:13px;">{section["description"]}</p>'
            )

        for field_data in section["fields"]:
            label = field_data["label"]
            field_type = field_data["field_type"]
            raw_value = field_data["raw_value"]
            display_value = field_data["display_value"]

            if field_type == "REPEATING_GROUP":
                # Render repeating group instances as a table
                instances = raw_value if isinstance(raw_value, list) else []
                if instances:
                    html_parts.append(f"<h3>{label}</h3>")
                    html_parts.append('<table class="data-table">')
                    # Header from first instance keys
                    if instances and isinstance(instances[0], dict):
                        html_parts.append("<tr>")
                        for key in instances[0].keys():
                            html_parts.append(f"<th>{key}</th>")
                        html_parts.append("</tr>")
                        for inst in instances:
                            html_parts.append("<tr>")
                            for key in instances[0].keys():
                                val = inst.get(key, "")
                                html_parts.append(f"<td>{val}</td>")
                            html_parts.append("</tr>")
                    html_parts.append("</table>")
                continue

            if field_type == "TABLE_GRID" and isinstance(raw_value, list):
                # Render table fields
                if raw_value:
                    html_parts.append(f"<h3>{label}</h3>")
                    html_parts.append('<table class="data-table">')
                    if isinstance(raw_value[0], dict):
                        html_parts.append("<tr>")
                        for key in raw_value[0].keys():
                            html_parts.append(f"<th>{key}</th>")
                        html_parts.append("</tr>")
                        for row in raw_value:
                            html_parts.append("<tr>")
                            for key in raw_value[0].keys():
                                html_parts.append(f"<td>{row.get(key, '')}</td>")
                            html_parts.append("</tr>")
                    html_parts.append("</table>")
                continue

            # Regular field
            html_parts.append('<div class="field-row">')
            html_parts.append(f'<div class="field-label">{label}</div>')
            if display_value:
                html_parts.append(f'<div class="field-value">{display_value}</div>')
            else:
                html_parts.append(
                    '<div class="field-value empty-value">Not provided</div>'
                )
            html_parts.append("</div>")

    # Evidence
    if data["evidence"]:
        html_parts.append('<hr class="section-divider">')
        html_parts.append("<h2>EVIDENCE</h2>")
        html_parts.append(
            '<table class="data-table"><tr><th>Type</th><th>File</th><th>Description</th><th>Verified</th></tr>'
        )
        for e in data["evidence"]:
            verified = "Yes" if e["verified"] else "No"
            html_parts.append(
                f'<tr><td>{e["type"]}</td><td>{e["filename"]}</td>'
                f'<td>{e["description"]}</td><td>{verified}</td></tr>'
            )
        html_parts.append("</table>")

    # Attachments
    if data["attachments"]:
        html_parts.append('<hr class="section-divider">')
        html_parts.append("<h2>ATTACHMENTS</h2>")
        html_parts.append(
            '<table class="data-table"><tr><th>File</th><th>Size</th><th>Description</th></tr>'
        )
        for a in data["attachments"]:
            html_parts.append(
                f'<tr><td>{a["filename"]}</td><td>{a["size"]}</td>'
                f'<td>{a["description"]}</td></tr>'
            )
        html_parts.append("</table>")

    # Notes
    if meta["notes"]:
        html_parts.append('<hr class="section-divider">')
        html_parts.append("<h2>NOTES</h2>")
        html_parts.append(f'<p>{meta["notes"]}</p>')

    # Footer
    html_parts.append(
        '<div class="footer">'
        "SITADC Youth Organization &middot; SITADC Youth Hub<br>"
        f"Report Reference: {meta['reference_number']}<br>"
        f"Generated: {timezone.now().strftime('%Y-%m-%d %H:%M:%S')}"
        "</div>"
    )
    html_parts.append("</body></html>")

    html_content = "\n".join(html_parts)
    safe_title = re.sub(r"[^\w\s\-]", "", meta["title"])[:60].strip()
    filename = f"{meta['reference_number']}-{safe_title}.html"

    response = HttpResponse(html_content, content_type="text/html")
    response["Content-Disposition"] = f'attachment; filename="{filename}"'
    return response


# ---------------------------------------------------------------------------
# CSV Export
# ---------------------------------------------------------------------------


def export_csv(report: Any) -> HttpResponse:
    """Export report as CSV with structured field data."""
    data = _build_structured_data(report)
    meta = data["meta"]

    output = io.StringIO()
    writer = csv.writer(output)

    # Header info
    writer.writerow(["Report Information", ""])
    writer.writerow(["Reference Number", meta["reference_number"]])
    writer.writerow(["Title", meta["title"]])
    writer.writerow(["Template", meta["template"]])
    writer.writerow(["Category", meta["category"]])
    writer.writerow(["Reporting Period", meta["reporting_period"]])
    writer.writerow(["Status", meta["status"]])
    writer.writerow(["Confidentiality", meta["confidentiality"]])
    writer.writerow(["Prepared By", meta["owner"]])
    writer.writerow(["Department", meta["department"]])
    writer.writerow(["Version", meta["version_number"]])
    writer.writerow([])

    # Sections
    for section in data["sections"]:
        if not section["fields"]:
            continue
        writer.writerow([section["name"].upper(), ""])
        for field_data in section["fields"]:
            if field_data["field_type"] == "TABLE_GRID":
                continue
            if field_data["field_type"] == "REPEATING_GROUP":
                continue
            display = field_data["display_value"] or "Not provided"
            writer.writerow(["", field_data["label"], display])
        writer.writerow([])

    # Evidence
    if data["evidence"]:
        writer.writerow(["Evidence", ""])
        writer.writerow(["Type", "File", "Description", "Verified"])
        for e in data["evidence"]:
            writer.writerow([e["type"], e["filename"], e["description"], e["verified"]])

    csv_content = output.getvalue()
    safe_title = re.sub(r"[^\w\s\-]", "", meta["title"])[:60].strip()
    filename = f"{meta['reference_number']}-{safe_title}.csv"

    response = HttpResponse(csv_content, content_type="text/csv")
    response["Content-Disposition"] = f'attachment; filename="{filename}"'
    return response


# ---------------------------------------------------------------------------
# JSON Export
# ---------------------------------------------------------------------------


def export_json(report: Any) -> HttpResponse:
    """Export report as JSON with structured data."""
    data = _build_structured_data(report)
    json_content = json.dumps(data, indent=2, default=str)

    meta = data["meta"]
    safe_title = re.sub(r"[^\w\s\-]", "", meta["title"])[:60].strip()
    filename = f"{meta['reference_number']}-{safe_title}.json"

    response = HttpResponse(json_content, content_type="application/json")
    response["Content-Disposition"] = f'attachment; filename="{filename}"'
    return response


# ---------------------------------------------------------------------------
# Excel Export (XLSX)
# ---------------------------------------------------------------------------


def export_xlsx(report: Any) -> HttpResponse:
    """Export report as XLSX with centered logo, structured sections, and styling."""
    try:
        from openpyxl import Workbook
        from openpyxl.drawing.image import Image as XLImage
        from openpyxl.styles import Alignment, Border, Font, PatternFill, Side
        from openpyxl.utils import get_column_letter
    except ImportError:
        return export_csv(report)

    data = _build_structured_data(report)
    meta = data["meta"]

    wb = Workbook()
    ws = wb.active
    ws.title = "Report"

    navy = "003366"
    light_grey = "EEF2F7"
    header_font = Font(bold=True, size=11, color="FFFFFF")
    header_fill = PatternFill(start_color=navy, end_color=navy, fill_type="solid")
    meta_label_font = Font(bold=True, size=10, color="444444")
    section_font = Font(bold=True, size=12, color=navy)
    field_label_font = Font(bold=True, size=10)
    thin_border = Border(
        left=Side(style="thin", color="CCCCCC"),
        right=Side(style="thin", color="CCCCCC"),
        top=Side(style="thin", color="CCCCCC"),
        bottom=Side(style="thin", color="CCCCCC"),
    )

    row_index = 1

    # Logo
    logo_path = _resolve_logo_path()
    if logo_path:
        try:
            img = XLImage(logo_path)
            img.width = 150
            img.height = 75
            ws.add_image(img, f"B{row_index}")
            row_index += 6  # Leave space for logo
        except Exception:
            pass

    # Title
    ws.merge_cells(f"A{row_index}:D{row_index}")
    title_cell = ws.cell(row=row_index, column=1, value=meta["title"].upper())
    title_cell.font = Font(bold=True, size=16, color=navy)
    title_cell.alignment = Alignment(horizontal="center")
    row_index += 2

    # Draft banner
    if meta["is_draft"]:
        ws.merge_cells(f"A{row_index}:D{row_index}")
        draft_cell = ws.cell(
            row=row_index, column=1, value="DRAFT — NOT FOR OFFICIAL USE"
        )
        draft_cell.font = Font(bold=True, size=11, color="B22222")
        draft_cell.alignment = Alignment(horizontal="center")
        row_index += 2

    # Metadata block
    meta_items = [
        ("Category", meta["category"]),
        ("Template", meta["template"]),
        ("Reference Number", meta["reference_number"]),
    ]
    if meta["reporting_period"]:
        meta_items.append(("Reporting Period", meta["reporting_period"]))
    if meta["template_version"]:
        meta_items.append(("Template Version", meta["template_version"]))
    if meta["department"]:
        meta_items.append(("Department", meta["department"]))
    meta_items.extend([
        ("Status", meta["status"]),
        ("Prepared By", meta["owner"]),
        ("Generated", timezone.now().strftime("%Y-%m-%d %H:%M")),
    ])

    for label, value in meta_items:
        ws.cell(row=row_index, column=1, value=label).font = meta_label_font
        ws.cell(row=row_index, column=1).fill = PatternFill(
            start_color=light_grey, end_color=light_grey, fill_type="solid"
        )
        ws.cell(row=row_index, column=1).border = thin_border
        ws.cell(row=row_index, column=2, value=str(value)).border = thin_border
        ws.merge_cells(f"B{row_index}:D{row_index}")
        row_index += 1

    row_index += 1

    # Sections
    for section in data["sections"]:
        if not section["fields"]:
            continue

        ws.merge_cells(f"A{row_index}:D{row_index}")
        ws.cell(
            row=row_index, column=1, value=section["name"].upper()
        ).font = section_font
        row_index += 1

        has_table = False
        for field_data in section["fields"]:
            ftype = field_data["field_type"]
            label = field_data["label"]
            raw_value = field_data["raw_value"]
            display_value = field_data["display_value"]

            if ftype == "TABLE_GRID" and isinstance(raw_value, list) and raw_value:
                has_table = True
                # Render as a sub-table
                if isinstance(raw_value[0], dict):
                    headers = list(raw_value[0].keys())
                    for col_idx, hdr in enumerate(headers, 1):
                        cell = ws.cell(row=row_index, column=col_idx, value=hdr)
                        cell.font = header_font
                        cell.fill = header_fill
                        cell.border = thin_border
                    row_index += 1
                    for row_data in raw_value:
                        for col_idx, hdr in enumerate(headers, 1):
                            cell = ws.cell(
                                row=row_index,
                                column=col_idx,
                                value=str(row_data.get(hdr, "")),
                            )
                            cell.border = thin_border
                        row_index += 1
                row_index += 1
            elif ftype == "REPEATING_GROUP" and isinstance(raw_value, list) and raw_value:
                has_table = True
                if raw_value and isinstance(raw_value[0], dict):
                    headers = list(raw_value[0].keys())
                    for col_idx, hdr in enumerate(headers, 1):
                        cell = ws.cell(row=row_index, column=col_idx, value=hdr)
                        cell.font = header_font
                        cell.fill = header_fill
                        cell.border = thin_border
                    row_index += 1
                    for inst in raw_value:
                        for col_idx, hdr in enumerate(headers, 1):
                            cell = ws.cell(
                                row=row_index,
                                column=col_idx,
                                value=str(inst.get(hdr, "")),
                            )
                            cell.border = thin_border
                        row_index += 1
                row_index += 1
            else:
                ws.cell(row=row_index, column=1, value=label).font = field_label_font
                value_text = display_value if display_value else "Not provided"
                ws.cell(row=row_index, column=2, value=str(value_text))
                ws.merge_cells(f"B{row_index}:D{row_index}")
                row_index += 1

        row_index += 1

    # Evidence
    if data["evidence"]:
        ws.cell(row=row_index, column=1, value="EVIDENCE").font = section_font
        row_index += 1
        ev_headers = ["Type", "File", "Description", "Verified"]
        for col_idx, hdr in enumerate(ev_headers, 1):
            cell = ws.cell(row=row_index, column=col_idx, value=hdr)
            cell.font = header_font
            cell.fill = header_fill
            cell.border = thin_border
        row_index += 1
        for ev in data["evidence"]:
            ws.cell(row=row_index, column=1, value=ev["type"]).border = thin_border
            ws.cell(row=row_index, column=2, value=ev["filename"]).border = thin_border
            ws.cell(row=row_index, column=3, value=ev["description"]).border = thin_border
            ws.cell(row=row_index, column=4, value="Yes" if ev["verified"] else "No").border = thin_border
            row_index += 1
        row_index += 1

    # Notes
    if meta["notes"]:
        ws.cell(row=row_index, column=1, value="NOTES").font = section_font
        row_index += 1
        ws.merge_cells(f"A{row_index}:D{row_index}")
        ws.cell(row=row_index, column=1, value=meta["notes"])
        row_index += 1

    # Adjust column widths
    ws.column_dimensions["A"].width = 25
    ws.column_dimensions["B"].width = 30
    ws.column_dimensions["C"].width = 25
    ws.column_dimensions["D"].width = 25

    # Footer row
    row_index += 2
    ws.merge_cells(f"A{row_index}:D{row_index}")
    ws.cell(
        row=row_index,
        column=1,
        value=f"SITADC Youth Hub — Generated {timezone.now().strftime('%Y-%m-%d %H:%M')}",
    ).font = Font(size=9, italic=True, color="666666")

    buffer = io.BytesIO()
    wb.save(buffer)
    buffer.seek(0)

    safe_title = re.sub(r"[^\w\s\-]", "", meta["title"])[:60].strip()
    filename = f"{meta['reference_number']}-{safe_title}.xlsx"

    response = HttpResponse(
        buffer,
        content_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
    )
    response["Content-Disposition"] = f'attachment; filename="{filename}"'
    return response


# ---------------------------------------------------------------------------
# PDF Export (using reportlab)
# ---------------------------------------------------------------------------


def export_pdf(report: Any) -> HttpResponse:
    """Export report as PDF with centered logo, structured sections, and professional layout."""
    try:
        from reportlab.lib import colors
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
        from reportlab.lib.units import mm
        from reportlab.platypus import (
            Image as RLImage,
            KeepTogether,
            PageBreak,
            Paragraph,
            SimpleDocTemplate,
            Spacer,
            Table,
            TableStyle,
        )
    except ImportError:
        return export_html(report)

    data = _build_structured_data(report)
    meta = data["meta"]

    buffer = io.BytesIO()
    doc = SimpleDocTemplate(
        buffer,
        pagesize=A4,
        leftMargin=18 * mm,
        rightMargin=18 * mm,
        topMargin=20 * mm,
        bottomMargin=20 * mm,
        title=meta["title"],
        author="SITADC Youth Hub",
    )

    styles = getSampleStyleSheet()
    page_width = A4[0] - 36 * mm  # Usable width

    # Custom styles
    title_style = ParagraphStyle(
        "ReportTitle",
        parent=styles["Title"],
        fontName="Helvetica-Bold",
        fontSize=16,
        leading=20,
        textColor=colors.HexColor("#003366"),
        alignment=1,  # center
        spaceAfter=6,
    )
    heading_style = ParagraphStyle(
        "SectionHeading",
        parent=styles["Heading2"],
        fontName="Helvetica-Bold",
        fontSize=12,
        leading=15,
        textColor=colors.HexColor("#003366"),
        spaceBefore=14,
        spaceAfter=6,
        borderWidth=0,
        borderPadding=0,
    )
    subheading_style = ParagraphStyle(
        "SubHeading",
        parent=styles["Heading3"],
        fontName="Helvetica-Bold",
        fontSize=10,
        leading=13,
        textColor=colors.HexColor("#333333"),
        spaceBefore=8,
        spaceAfter=4,
    )
    label_style = ParagraphStyle(
        "FieldLabel",
        parent=styles["Normal"],
        fontName="Helvetica-Bold",
        fontSize=9,
        leading=12,
        textColor=colors.HexColor("#444444"),
    )
    value_style = ParagraphStyle(
        "FieldValue",
        parent=styles["Normal"],
        fontName="Helvetica",
        fontSize=9,
        leading=12,
        textColor=colors.HexColor("#333333"),
    )
    empty_style = ParagraphStyle(
        "EmptyValue",
        parent=styles["Normal"],
        fontName="Helvetica-Oblique",
        fontSize=9,
        leading=12,
        textColor=colors.HexColor("#999999"),
    )
    meta_key_style = ParagraphStyle(
        "MetaKey",
        parent=styles["Normal"],
        fontName="Helvetica-Bold",
        fontSize=9,
        leading=11,
        textColor=colors.HexColor("#333333"),
    )
    meta_val_style = ParagraphStyle(
        "MetaVal",
        parent=styles["Normal"],
        fontName="Helvetica",
        fontSize=9,
        leading=11,
        textColor=colors.HexColor("#333333"),
    )
    footer_style = ParagraphStyle(
        "Footer",
        parent=styles["Normal"],
        fontName="Helvetica",
        fontSize=8,
        leading=10,
        textColor=colors.HexColor("#666666"),
        alignment=1,  # center
    )

    elements = []

    # ---- Centered Logo ----
    logo_path = _resolve_logo_path()
    if logo_path:
        try:
            from reportlab.lib.utils import ImageReader

            img_reader = ImageReader(logo_path)
            iw, ih = img_reader.getSize()
            aspect = ih / iw
            logo_width = 50 * mm
            logo_height = logo_width * aspect
            if logo_height > 35 * mm:
                logo_height = 35 * mm
                logo_width = logo_height / aspect

            logo_img = RLImage(logo_path, width=logo_width, height=logo_height)
            # Wrap in a table to center it
            logo_table = Table([[logo_img]], colWidths=[page_width])
            logo_table.setStyle(
                TableStyle([("ALIGN", (0, 0), (-1, -1), "CENTER")])
            )
            elements.append(logo_table)
            elements.append(Spacer(1, 8))
        except Exception as e:
            logger.warning("Failed to load logo for PDF: %s", e)

    # ---- Draft Banner ----
    if meta["is_draft"]:
        draft_style = ParagraphStyle(
            "DraftBanner",
            parent=styles["Normal"],
            fontName="Helvetica-Bold",
            fontSize=10,
            textColor=colors.white,
            alignment=1,
            backColor=colors.HexColor("#B22222"),
            borderPadding=6,
        )
        elements.append(
            KeepTogether(
                [Paragraph("DRAFT — NOT FOR OFFICIAL USE", draft_style)]
            )
        )
        elements.append(Spacer(1, 10))

    # ---- Report Title (centered) ----
    elements.append(Paragraph(meta["title"].upper(), title_style))
    elements.append(Spacer(1, 6))

    # ---- Metadata Table ----
    meta_rows = [
        ("Category", meta["category"]),
        ("Template", meta["template"]),
        ("Reference", meta["reference_number"]),
    ]
    if meta["reporting_period"]:
        meta_rows.append(("Reporting Period", meta["reporting_period"]))
    if meta["template_version"]:
        meta_rows.append(("Template Version", meta["template_version"]))
    if meta["department"]:
        meta_rows.append(("Department", meta["department"]))
    meta_rows.append(("Status", meta["status"]))
    if meta["owner"]:
        meta_rows.append(("Prepared By", meta["owner"]))
    meta_rows.append(("Generated", timezone.now().strftime("%Y-%m-%d %H:%M")))

    meta_table_data = [
        [Paragraph(key, meta_key_style), Paragraph(str(val), meta_val_style)]
        for key, val in meta_rows
    ]
    meta_table = Table(meta_table_data, colWidths=[45 * mm, page_width - 45 * mm])
    meta_table.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (0, -1), colors.HexColor("#EEF2F7")),
                ("FONTNAME", (0, 0), (0, -1), "Helvetica-Bold"),
                ("GRID", (0, 0), (-1, -1), 0.4, colors.HexColor("#CCCCCC")),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("TOPPADDING", (0, 0), (-1, -1), 3),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
                ("LEFTPADDING", (0, 0), (-1, -1), 6),
            ]
        )
    )
    elements.append(meta_table)
    elements.append(Spacer(1, 12))

    # ---- Horizontal Line ----
    line_table = Table([[""]], colWidths=[page_width])
    line_table.setStyle(
        TableStyle(
            [
                ("LINEBELOW", (0, 0), (-1, 0), 1, colors.HexColor("#003366")),
            ]
        )
    )
    elements.append(line_table)
    elements.append(Spacer(1, 8))

    # ---- Report Sections ----
    for section in data["sections"]:
        if not section["fields"]:
            continue

        section_elements = []
        section_elements.append(
            Paragraph(section["name"].upper(), heading_style)
        )

        if section["description"]:
            section_elements.append(
                Paragraph(section["description"], ParagraphStyle(
                    "SectionDesc",
                    parent=styles["Normal"],
                    fontName="Helvetica-Oblique",
                    fontSize=9,
                    leading=12,
                    textColor=colors.HexColor("#666666"),
                    spaceAfter=4,
                ))
            )

        for field_data in section["fields"]:
            ftype = field_data["field_type"]
            label = field_data["label"]
            raw_value = field_data["raw_value"]
            display_value = field_data["display_value"]

            if ftype == "TABLE_GRID" and isinstance(raw_value, list) and raw_value:
                # Render as a reportlab table
                if isinstance(raw_value[0], dict):
                    section_elements.append(Paragraph(label, subheading_style))
                    headers = list(raw_value[0].keys())
                    table_data = [headers]
                    for row_data in raw_value:
                        table_data.append([str(row_data.get(h, "")) for h in headers])

                    col_count = len(headers)
                    col_width = page_width / col_count
                    tbl = Table(
                        table_data,
                        colWidths=[col_width] * col_count,
                        repeatRows=1,
                    )
                    tbl.setStyle(
                        TableStyle(
                            [
                                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#003366")),
                                ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                                ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                                ("FONTSIZE", (0, 0), (-1, -1), 8),
                                ("GRID", (0, 0), (-1, -1), 0.25, colors.HexColor("#CCCCCC")),
                                ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#EEF2F7")]),
                                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                                ("TOPPADDING", (0, 0), (-1, -1), 3),
                                ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
                            ]
                        )
                    )
                    section_elements.append(tbl)
                    section_elements.append(Spacer(1, 6))
            elif ftype == "REPEATING_GROUP" and isinstance(raw_value, list) and raw_value:
                section_elements.append(Paragraph(label, subheading_style))
                if raw_value and isinstance(raw_value[0], dict):
                    headers = list(raw_value[0].keys())
                    table_data = [headers]
                    for inst in raw_value:
                        table_data.append([str(inst.get(h, "")) for h in headers])
                    col_count = len(headers)
                    col_width = page_width / col_count
                    tbl = Table(table_data, colWidths=[col_width] * col_count, repeatRows=1)
                    tbl.setStyle(
                        TableStyle([
                            ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#003366")),
                            ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                            ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                            ("FONTSIZE", (0, 0), (-1, -1), 8),
                            ("GRID", (0, 0), (-1, -1), 0.25, colors.HexColor("#CCCCCC")),
                            ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#EEF2F7")]),
                            ("VALIGN", (0, 0), (-1, -1), "TOP"),
                            ("TOPPADDING", (0, 0), (-1, -1), 3),
                            ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
                        ])
                    )
                    section_elements.append(tbl)
                    section_elements.append(Spacer(1, 6))
            else:
                # Regular field — label + value
                val_text = display_value if display_value else ""
                val_para_style = value_style if val_text else empty_style
                val_display = val_text if val_text else "Not provided"

                field_table = Table(
                    [
                        [Paragraph(label, label_style)],
                        [Paragraph(val_display, val_para_style)],
                    ],
                    colWidths=[page_width],
                )
                field_table.setStyle(
                    TableStyle(
                        [
                            ("TOPPADDING", (0, 0), (-1, -1), 2),
                            ("BOTTOMPADDING", (0, 0), (-1, -1), 2),
                            ("LINEBELOW", (0, -1), (-1, -1), 0.3, colors.HexColor("#EEEEEE")),
                        ]
                    )
                )
                section_elements.append(field_table)

        elements.append(KeepTogether(section_elements[:3]))  # Keep at least section heading together
        for elem in section_elements[3:]:
            elements.append(elem)
        elements.append(Spacer(1, 6))

    # ---- Evidence ----
    if data["evidence"]:
        elements.append(Paragraph("EVIDENCE", heading_style))
        ev_headers = ["Type", "File", "Description", "Verified"]
        ev_data = [ev_headers]
        for ev in data["evidence"]:
            ev_data.append([
                ev["type"],
                ev["filename"],
                ev["description"],
                "Yes" if ev["verified"] else "No",
            ])
        col_count = 4
        col_width = page_width / col_count
        ev_table = Table(ev_data, colWidths=[col_width] * col_count, repeatRows=1)
        ev_table.setStyle(
            TableStyle([
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#003366")),
                ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                ("FONTSIZE", (0, 0), (-1, -1), 8),
                ("GRID", (0, 0), (-1, -1), 0.25, colors.HexColor("#CCCCCC")),
                ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#EEF2F7")]),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("TOPPADDING", (0, 0), (-1, -1), 3),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
            ])
        )
        elements.append(ev_table)
        elements.append(Spacer(1, 10))

    # ---- Notes ----
    if meta["notes"]:
        elements.append(Paragraph("NOTES", heading_style))
        elements.append(Paragraph(meta["notes"], value_style))
        elements.append(Spacer(1, 10))

    # ---- Footer ----
    elements.append(Spacer(1, 20))
    line_table2 = Table([[""]], colWidths=[page_width])
    line_table2.setStyle(
        TableStyle([
            ("LINEABOVE", (0, 0), (-1, 0), 0.5, colors.HexColor("#CCCCCC")),
        ])
    )
    elements.append(line_table2)
    elements.append(Spacer(1, 4))
    elements.append(
        Paragraph(
            "SITADC Youth Organization &middot; SITADC Youth Hub",
            footer_style,
        )
    )
    elements.append(
        Paragraph(
            f"Reference: {meta['reference_number']} &middot; "
            f"Generated: {timezone.now().strftime('%Y-%m-%d %H:%M:%S')}",
            footer_style,
        )
    )

    # Build PDF
    doc.build(elements)
    buffer.seek(0)

    safe_title = re.sub(r"[^\w\s\-]", "", meta["title"])[:60].strip()
    filename = f"{meta['reference_number']}-{safe_title}.pdf"

    response = HttpResponse(buffer, content_type="application/pdf")
    response["Content-Disposition"] = f'attachment; filename="{filename}"'
    return response


# ---------------------------------------------------------------------------
# DOCX Export (using python-docx)
# ---------------------------------------------------------------------------


def export_docx(report: Any) -> HttpResponse:
    """Export report as Word DOCX with centered logo, structured sections, and tables."""
    try:
        from docx import Document
        from docx.enum.section import WD_ORIENT
        from docx.enum.table import WD_TABLE_ALIGNMENT
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
        from docx.shared import Cm, Inches, Pt, RGBColor
    except ImportError:
        return export_html(report)

    data = _build_structured_data(report)
    meta = data["meta"]

    document = Document()

    # Page setup
    section = document.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    # Default style
    style = document.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)

    navy = RGBColor(0x00, 0x33, 0x66)
    grey = RGBColor(0x55, 0x55, 0x55)

    def _shade_cell(cell, hex_color):
        from docx.oxml import OxmlElement
        tc_pr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), hex_color)
        tc_pr.append(shd)

    def _add_meta_row(label, value):
        meta_table = document.add_table(rows=1, cols=2)
        meta_table.style = "Table Grid"
        meta_table.alignment = WD_TABLE_ALIGNMENT.LEFT
        cells = meta_table.rows[0].cells
        cells[0].text = label
        cells[1].text = str(value)
        for paragraph in cells[0].paragraphs:
            for r in paragraph.runs:
                r.bold = True
                r.font.size = Pt(9)
        for paragraph in cells[1].paragraphs:
            for r in paragraph.runs:
                r.font.size = Pt(9)
        _shade_cell(cells[0], "EEF2F7")
        return meta_table

    # ---- Centered Logo ----
    logo_path = _resolve_logo_path()
    if logo_path:
        try:
            from docx.shared import Emu

            p = document.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run()
            run.add_picture(logo_path, width=Inches(2.0))
        except Exception as e:
            logger.warning("Failed to load logo for DOCX: %s", e)

    # ---- Draft Banner ----
    if meta["is_draft"]:
        draft = document.add_paragraph()
        draft.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = draft.add_run("DRAFT — NOT FOR OFFICIAL USE")
        run.bold = True
        run.font.color.rgb = RGBColor(0xB2, 0x22, 0x22)
        run.font.size = Pt(11)

    # ---- Title (centered) ----
    title_para = document.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_para.add_run(meta["title"].upper())
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = navy

    document.add_paragraph()  # Spacer

    # ---- Metadata ----
    meta_items = [
        ("Category", meta["category"]),
        ("Template", meta["template"]),
        ("Reference Number", meta["reference_number"]),
    ]
    if meta["reporting_period"]:
        meta_items.append(("Reporting Period", meta["reporting_period"]))
    if meta["template_version"]:
        meta_items.append(("Template Version", meta["template_version"]))
    if meta["department"]:
        meta_items.append(("Department", meta["department"]))
    meta_items.extend([
        ("Status", meta["status"]),
        ("Prepared By", meta["owner"]),
        ("Generated", timezone.now().strftime("%Y-%m-%d %H:%M")),
    ])

    for label, value in meta_items:
        _add_meta_row(label, value)

    document.add_paragraph()  # Spacer

    # ---- Sections ----
    for section_data in data["sections"]:
        if not section_data["fields"]:
            continue

        # Section heading
        heading = document.add_heading(level=2)
        run = heading.add_run(section_data["name"].upper())
        run.font.color.rgb = navy

        if section_data["description"]:
            desc = document.add_paragraph(section_data["description"])
            desc.runs[0].font.italic = True
            desc.runs[0].font.color.rgb = grey
            desc.runs[0].font.size = Pt(9)

        for field_data in section_data["fields"]:
            ftype = field_data["field_type"]
            label = field_data["label"]
            raw_value = field_data["raw_value"]
            display_value = field_data["display_value"]

            if ftype == "TABLE_GRID" and isinstance(raw_value, list) and raw_value:
                if isinstance(raw_value[0], dict):
                    # Add table heading
                    sub = document.add_paragraph()
                    sub_run = sub.add_run(label)
                    sub_run.bold = True
                    sub_run.font.size = Pt(10)

                    headers = list(raw_value[0].keys())
                    table = document.add_table(rows=1, cols=len(headers))
                    table.style = "Table Grid"
                    table.alignment = WD_TABLE_ALIGNMENT.LEFT

                    # Header row
                    for col_idx, hdr in enumerate(headers):
                        cell = table.rows[0].cells[col_idx]
                        cell.text = str(hdr)
                        for paragraph in cell.paragraphs:
                            for r in paragraph.runs:
                                r.bold = True
                                r.font.size = Pt(9)
                                r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                        _shade_cell(cell, "003366")

                    # Data rows
                    for row_data in raw_value:
                        cells = table.add_row().cells
                        for col_idx, hdr in enumerate(headers):
                            cells[col_idx].text = str(row_data.get(hdr, ""))
                            for paragraph in cells[col_idx].paragraphs:
                                for r in paragraph.runs:
                                    r.font.size = Pt(9)
            elif ftype == "REPEATING_GROUP" and isinstance(raw_value, list) and raw_value:
                sub = document.add_paragraph()
                sub_run = sub.add_run(label)
                sub_run.bold = True
                sub_run.font.size = Pt(10)

                if raw_value and isinstance(raw_value[0], dict):
                    headers = list(raw_value[0].keys())
                    table = document.add_table(rows=1, cols=len(headers))
                    table.style = "Table Grid"
                    for col_idx, hdr in enumerate(headers):
                        cell = table.rows[0].cells[col_idx]
                        cell.text = str(hdr)
                        for paragraph in cell.paragraphs:
                            for r in paragraph.runs:
                                r.bold = True
                                r.font.size = Pt(9)
                                r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                        _shade_cell(cell, "003366")
                    for inst in raw_value:
                        cells = table.add_row().cells
                        for col_idx, hdr in enumerate(headers):
                            cells[col_idx].text = str(inst.get(hdr, ""))
                            for paragraph in cells[col_idx].paragraphs:
                                for r in paragraph.runs:
                                    r.font.size = Pt(9)
            else:
                # Regular field
                p = document.add_paragraph()
                label_run = p.add_run(f"{label}: ")
                label_run.bold = True
                label_run.font.size = Pt(10)
                val_text = display_value if display_value else "Not provided"
                val_run = p.add_run(val_text)
                val_run.font.size = Pt(10)
                if not display_value:
                    val_run.font.italic = True
                    val_run.font.color.rgb = grey

    # ---- Evidence ----
    if data["evidence"]:
        heading = document.add_heading(level=2)
        run = heading.add_run("EVIDENCE")
        run.font.color.rgb = navy

        table = document.add_table(rows=1, cols=4)
        table.style = "Table Grid"
        headers = ["Type", "File", "Description", "Verified"]
        for col_idx, hdr in enumerate(headers):
            cell = table.rows[0].cells[col_idx]
            cell.text = hdr
            for paragraph in cell.paragraphs:
                for r in paragraph.runs:
                    r.bold = True
                    r.font.size = Pt(9)
                    r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            _shade_cell(cell, "003366")

        for ev in data["evidence"]:
            cells = table.add_row().cells
            cells[0].text = ev["type"]
            cells[1].text = ev["filename"]
            cells[2].text = ev["description"]
            cells[3].text = "Yes" if ev["verified"] else "No"
            for cell in cells:
                for paragraph in cell.paragraphs:
                    for r in paragraph.runs:
                        r.font.size = Pt(9)

    # ---- Notes ----
    if meta["notes"]:
        heading = document.add_heading(level=2)
        run = heading.add_run("NOTES")
        run.font.color.rgb = navy
        document.add_paragraph(meta["notes"])

    # ---- Footer ----
    document.add_paragraph()
    footer = document.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run(
        f"SITADC Youth Organization · SITADC Youth Hub\n"
        f"Reference: {meta['reference_number']} · "
        f"Generated: {timezone.now().strftime('%Y-%m-%d %H:%M:%S')}"
    )
    run.font.size = Pt(8)
    run.font.color.rgb = grey

    # Save
    buffer = io.BytesIO()
    document.save(buffer)
    buffer.seek(0)

    safe_title = re.sub(r"[^\w\s\-]", "", meta["title"])[:60].strip()
    filename = f"{meta['reference_number']}-{safe_title}.docx"

    response = HttpResponse(
        buffer,
        content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )
    response["Content-Disposition"] = f'attachment; filename="{filename}"'
    return response
