"""Multi-format export adapters for Organizational Registers.

CSV/XLSX/DOCX/PDF adapters mirror ``apps.meal.exports`` and
``apps.programs.report_exports``.  Every export is permission checked and
records an immutable activity timeline entry.
"""

from __future__ import annotations

import csv
import json
import logging
from io import BytesIO, StringIO
from pathlib import Path

from django.conf import settings
from django.core.exceptions import PermissionDenied
from django.http import HttpResponse

from apps.rbac.authorization import user_has_permission

from .models import RegisterActivity, RegisterEntry
from .permissions import REGISTER_EXPORT, REGISTER_MANAGE
from .selectors import visible_entries

logger = logging.getLogger(__name__)

LOGO_PATH = Path(settings.BASE_DIR) / "static" / "images" / "app_logo.png"


def formula_safe_csv_value(value) -> str:
    text = str(value or "")
    if text.startswith(("=", "+", "-", "@", "\t", "\r")):
        return f"'{text}"
    return text


def _require_export_permission(user) -> None:
    if not (
        user_has_permission(user, REGISTER_EXPORT)
        or user_has_permission(user, REGISTER_MANAGE)
    ):
        raise PermissionDenied("Register export permission is required.")


def _csv_response(user, headers, rows, filename: str, event: str) -> HttpResponse:
    output = StringIO(newline="")
    writer = csv.writer(output)
    writer.writerow(headers)
    for row in rows:
        writer.writerow([formula_safe_csv_value(value) for value in row])
    response = HttpResponse(output.getvalue(), content_type="text/csv; charset=utf-8")
    response["Content-Disposition"] = f'attachment; filename="{filename}"'
    response["Cache-Control"] = "private, no-store"
    response["Pragma"] = "no-cache"
    response["X-Content-Type-Options"] = "nosniff"
    logger.info(
        event,
        extra={
            "register_event": {
                "action": event,
                "format": "csv",
                "actor_id": str(user.pk),
                "row_count": len(rows),
                "scope": "visible_registers",
            }
        },
    )
    return response


def _xlsx_styled_header(sheet, headers) -> None:
    from openpyxl.styles import Font, PatternFill  # type: ignore[import-untyped]

    fill = PatternFill(start_color="003366", end_color="003366", fill_type="solid")
    for col_idx, header in enumerate(headers, 1):
        cell = sheet.cell(row=1, column=col_idx, value=header)
        cell.font = Font(bold=True, color="FFFFFF")
        cell.fill = fill


def _xlsx_set_widths(sheet, widths: list) -> None:
    from openpyxl.utils import get_column_letter  # type: ignore[import-untyped]

    for col_idx, width in enumerate(widths, 1):
        sheet.column_dimensions[get_column_letter(col_idx)].width = width


def _build_xlsx(title: str, rows: list[list], headers: list, widths: list) -> BytesIO:
    from openpyxl import Workbook  # type: ignore[import-untyped]

    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "Register"
    _xlsx_styled_header(sheet, headers)
    for row in rows:
        sheet.append(row)
    _xlsx_set_widths(sheet, widths)
    buffer = BytesIO()
    workbook.save(buffer)
    buffer.seek(0)
    return buffer


def _build_docx(title: str, rows: list[list], headers: list) -> BytesIO:
    from docx import Document

    document = Document()
    document.add_heading(title, level=0)
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Light Grid Accent 1"
    for col_idx, header in enumerate(headers):
        table.rows[0].cells[col_idx].text = str(header)
    for row in rows:
        cells = table.add_row().cells
        for col_idx, value in enumerate(row):
            cells[col_idx].text = str(value)
    buffer = BytesIO()
    document.save(buffer)
    buffer.seek(0)
    return buffer


def _build_pdf(title: str, rows: list[list], headers: list) -> BytesIO:
    from reportlab.lib import colors  # type: ignore[import-untyped]
    from reportlab.lib.pagesizes import A4, landscape  # type: ignore[import-untyped]
    from reportlab.lib.styles import getSampleStyleSheet  # type: ignore[import-untyped]
    from reportlab.platypus import (  # type: ignore[import-untyped]
        SimpleDocTemplate,
        Table,
        TableStyle,
    )

    buffer = BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=landscape(A4), title=title)
    styles = getSampleStyleSheet()
    elements = [_reportlab_heading(title, styles)]
    data = [headers] + [[str(value) for value in row] for row in rows]
    table = Table(data, repeatRows=1)
    table.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#003366")),
                ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                ("FONTSIZE", (0, 0), (-1, 0), 8),
                ("FONTSIZE", (0, 1), (-1, -1), 7),
                ("GRID", (0, 0), (-1, -1), 0.25, colors.grey),
                (
                    "ROWBACKGROUNDS",
                    (0, 1),
                    (-1, -1),
                    [colors.white, colors.HexColor("#EEF2F7")],
                ),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
            ]
        )
    )
    elements.append(table)
    doc.build(elements)
    buffer.seek(0)
    return buffer


def _reportlab_heading(title: str, styles) -> object:
    from reportlab.lib import colors
    from reportlab.lib.styles import ParagraphStyle
    from reportlab.platypus import Paragraph

    style = ParagraphStyle(
        "RegisterTitle",
        parent=styles["Title"],
        fontName="Helvetica-Bold",
        fontSize=14,
        leading=18,
        textColor=colors.HexColor("#003366"),
    )
    return Paragraph(title, style)


REGISTER_HEADERS = [
    "Reference Number",
    "Title",
    "Register",
    "Category",
    "Confidentiality",
    "Approval Status",
    "Status",
    "Owner",
    "Directorate",
    "Program",
    "Project",
    "Reporting Period",
    "Created At",
    "Updated At",
]


def _entry_row(entry: RegisterEntry) -> list:
    created_at = entry.created_at
    if created_at is not None and getattr(created_at, "tzinfo", None) is not None:
        created_at = created_at.replace(tzinfo=None)
    updated_at = entry.updated_at
    if updated_at is not None and getattr(updated_at, "tzinfo", None) is not None:
        updated_at = updated_at.replace(tzinfo=None)
    return [
        entry.reference_number,
        entry.title,
        entry.register.name,
        entry.register.category.name,
        entry.get_confidentiality_display(),
        entry.get_approval_status_display(),
        entry.get_status_display(),
        entry.owner.full_name if entry.owner else "",
        entry.directorate.name if entry.directorate else "",
        entry.program.title if entry.program else "",
        entry.project.title if entry.project else "",
        (
            f"{entry.reporting_period_start} to {entry.reporting_period_end}"
            if entry.reporting_period_start or entry.reporting_period_end
            else ""
        ),
        created_at,
        updated_at,
    ]


def _record_export_activity(user, register=None, fmt: str = "csv") -> None:
    RegisterActivity.objects.create(
        register=register,
        action="EXPORTED",
        actor=user,
        new_status=fmt.upper(),
        comment=f"Register data exported as {fmt.upper()}.",
    )


def register_register_csv_response(user, register=None) -> HttpResponse:
    """Export the visible register entries as CSV."""
    _require_export_permission(user)
    queryset = visible_entries(user)
    if register is not None:
        queryset = queryset.filter(register=register)
    rows = [
        _entry_row(entry)
        for entry in queryset.select_related(
            "register__category", "owner", "directorate", "program", "project"
        )
    ]
    _record_export_activity(user, register=register, fmt="csv")
    return _csv_response(
        user,
        REGISTER_HEADERS,
        rows,
        "registers_export.csv",
        "registers_exported",
    )


def register_export_response(user, register=None, fmt: str = "csv") -> HttpResponse:
    """Export register entries in the requested format."""
    _require_export_permission(user)
    queryset = visible_entries(user)
    if register is not None:
        queryset = queryset.filter(register=register)
    entries = list(
        queryset.select_related(
            "register__category", "owner", "directorate", "program", "project"
        )
    )
    title = "Organizational Registers Export"
    rows = [_entry_row(entry) for entry in entries]
    widths = [18, 30, 22, 18, 16, 18, 14, 22, 18, 22, 22, 30, 22, 22]

    if fmt == "csv":
        _record_export_activity(user, register=register, fmt="csv")
        return _csv_response(
            user,
            REGISTER_HEADERS,
            rows,
            "registers_export.csv",
            "registers_exported",
        )
    if fmt == "json":
        _record_export_activity(user, register=register, fmt="json")
        safe_rows = [[formula_safe_csv_value(v) for v in row] for row in rows]
        payload = [dict(zip(REGISTER_HEADERS, row, strict=False)) for row in safe_rows]
        response = HttpResponse(
            json.dumps(payload, indent=2, default=str),
            content_type="application/json; charset=utf-8",
        )
        response["Content-Disposition"] = 'attachment; filename="registers_export.json"'
        response["Cache-Control"] = "private, no-store"
        response["X-Content-Type-Options"] = "nosniff"
        return response

    if fmt == "xlsx":
        _record_export_activity(user, register=register, fmt="xlsx")
        buffer = _build_xlsx(title, rows, REGISTER_HEADERS, widths)
        content_type = (
            "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
        )
        filename = "registers_export.xlsx"
    elif fmt == "docx":
        _record_export_activity(user, register=register, fmt="docx")
        buffer = _build_docx(title, rows, REGISTER_HEADERS)
        content_type = (
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
        filename = "registers_export.docx"
    elif fmt == "pdf":
        _record_export_activity(user, register=register, fmt="pdf")
        buffer = _build_pdf(title, rows, REGISTER_HEADERS)
        content_type = "application/pdf"
        filename = "registers_export.pdf"
    else:
        raise PermissionDenied(f"Unsupported export format: {fmt}")

    response = HttpResponse(buffer.getvalue(), content_type=content_type)
    response["Content-Disposition"] = f'attachment; filename="{filename}"'
    response["Cache-Control"] = "private, no-store"
    response["Pragma"] = "no-cache"
    response["X-Content-Type-Options"] = "nosniff"
    return response
