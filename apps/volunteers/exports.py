"""
Export helpers for volunteer register reporting (CSV, XLSX, DOCX, PDF).
"""

from __future__ import annotations

from io import BytesIO

from django.http import HttpResponse
from django.utils.translation import gettext as _


def _safe_cell(value) -> str:
    text = str(value or "")
    if text.startswith(("=", "+", "-", "@", "\t", "\r")):
        return f"'{text}"
    return text


def build_volunteer_rows(profiles, include_confidential: bool):
    """Yield normalized rows for the volunteer register export."""
    headers = [
        _("Reference Number"),
        _("Full Name"),
        _("Category"),
        _("Status"),
        _("Region"),
        _("District"),
    ]
    if include_confidential:
        headers.extend([_("Email"), _("Phone")])
    yield headers
    for profile in profiles.iterator(chunk_size=500):
        row = [
            profile.reference_number,
            profile.user.full_name,
            profile.category.name if profile.category_id else "",
            profile.get_status_display(),
            profile.region,
            profile.district,
        ]
        if include_confidential:
            row.extend([profile.email, profile.phone_number])
        yield [_safe_cell(value) for value in row]


def export_volunteer_csv(profiles, include_confidential: bool) -> HttpResponse:
    import csv

    response = HttpResponse(content_type="text/csv; charset=utf-8")
    response["Content-Disposition"] = 'attachment; filename="volunteer_register.csv"'
    response["Cache-Control"] = "private, no-store"
    response["X-Content-Type-Options"] = "nosniff"
    writer = csv.writer(response)
    row_count = 0
    for row in build_volunteer_rows(profiles, include_confidential):
        writer.writerow(row)
        row_count += 1
    response["X-Row-Count"] = str(row_count)
    return response


def export_volunteer_xlsx(profiles, include_confidential: bool) -> HttpResponse:
    from openpyxl import Workbook  # type: ignore[import-untyped]

    workbook = Workbook()
    sheet = workbook.active
    sheet.title = _("Volunteer Register")
    for row in build_volunteer_rows(profiles, include_confidential):
        sheet.append(row)
    for column in sheet.columns:
        width = max(
            (len(str(cell.value or "")) for cell in column),
            default=8,
        )
        sheet.column_dimensions[column[0].column_letter].width = min(width + 4, 60)
    output = BytesIO()
    workbook.save(output)
    response = HttpResponse(
        output.getvalue(),
        content_type=(
            "application/vnd.openxmlformats-officedocument." "spreadsheetml.sheet"
        ),
    )
    response["Content-Disposition"] = 'attachment; filename="volunteer_register.xlsx"'
    response["Cache-Control"] = "private, no-store"
    response["X-Content-Type-Options"] = "nosniff"
    return response


def export_volunteer_docx(profiles, include_confidential: bool) -> HttpResponse:
    from docx import Document

    document = Document()
    document.add_heading(_("Volunteer Register"), level=1)
    rows = list(build_volunteer_rows(profiles, include_confidential))
    if not rows:
        return _empty_docx()
    headers = rows[0]
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Light Grid Accent 1"
    for index, header in enumerate(headers):
        table.rows[0].cells[index].text = header
    for data_row in rows[1:]:
        cells = table.add_row().cells
        for index, value in enumerate(data_row):
            cells[index].text = value
    output = BytesIO()
    document.save(output)
    response = HttpResponse(
        output.getvalue(),
        content_type=(
            "application/vnd.openxmlformats-officedocument." "wordprocessingml.document"
        ),
    )
    response["Content-Disposition"] = 'attachment; filename="volunteer_register.docx"'
    response["Cache-Control"] = "private, no-store"
    response["X-Content-Type-Options"] = "nosniff"
    return response


def _empty_docx() -> HttpResponse:
    from docx import Document

    document = Document()
    document.add_paragraph(_("No volunteer records available."))
    output = BytesIO()
    document.save(output)
    response = HttpResponse(
        output.getvalue(),
        content_type=(
            "application/vnd.openxmlformats-officedocument." "wordprocessingml.document"
        ),
    )
    response["Content-Disposition"] = 'attachment; filename="volunteer_register.docx"'
    response["Cache-Control"] = "private, no-store"
    return response


def export_volunteer_pdf(profiles, include_confidential: bool) -> HttpResponse:
    from reportlab.lib import colors  # type: ignore[import-untyped]
    from reportlab.lib.pagesizes import A4, landscape  # type: ignore[import-untyped]
    from reportlab.lib.units import mm  # type: ignore[import-untyped]
    from reportlab.platypus import (  # type: ignore[import-untyped]
        SimpleDocTemplate,
        Spacer,
        Table,
        TableStyle,
    )

    rows = list(build_volunteer_rows(profiles, include_confidential))
    data = [list(header_row) for header_row in rows] or [[""]]
    output = BytesIO()
    document = SimpleDocTemplate(
        output,
        pagesize=landscape(A4),
        leftMargin=12 * mm,
        rightMargin=12 * mm,
        topMargin=12 * mm,
        bottomMargin=12 * mm,
        title=_("Volunteer Register"),
    )
    table = Table(data)
    table.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#0d6efd")),
                ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                ("FONTSIZE", (0, 0), (-1, -1), 8),
                ("GRID", (0, 0), (-1, -1), 0.4, colors.grey),
                (
                    "ROWBACKGROUNDS",
                    (0, 1),
                    (-1, -1),
                    [colors.white, colors.HexColor("#f1f5f9")],
                ),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
            ]
        )
    )
    document.build([Spacer(1, 6 * mm), table])
    response = HttpResponse(output.getvalue(), content_type="application/pdf")
    response["Content-Disposition"] = 'attachment; filename="volunteer_register.pdf"'
    response["Cache-Control"] = "private, no-store"
    response["X-Content-Type-Options"] = "nosniff"
    return response
