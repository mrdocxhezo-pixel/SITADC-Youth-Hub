"""Multi-format (XLSX, DOCX, PDF) report export adapters for programs."""

from __future__ import annotations

import logging
from io import BytesIO
from pathlib import Path

from django.conf import settings
from django.core.exceptions import PermissionDenied
from django.http import HttpResponse

from apps.rbac.authorization import user_has_permission

from .permissions import (
    PROGRAMMES_EXPORT,
    PROGRAMMES_MANAGE,
    PROJECTS_EXPORT,
    PROJECTS_MANAGE,
)
from .selectors import visible_programs, visible_projects

logger = logging.getLogger(__name__)

LOGO_PATH = Path(settings.BASE_DIR) / "static" / "images" / "app_logo.png"


# ── openpyxl helpers ──────────────────────────────────────────────────────


def _xlsx_styled_header(ws, headers: list) -> None:
    from openpyxl.styles import (  # type: ignore[import-untyped]
        Alignment,
        Border,
        Font,
        PatternFill,
        Side,
    )

    fill = PatternFill(start_color="003366", end_color="003366", fill_type="solid")
    font = Font(color="FFFFFF", bold=True, size=11)
    border = Border(
        left=Side(style="thin"),
        right=Side(style="thin"),
        top=Side(style="thin"),
        bottom=Side(style="thin"),
    )
    for col_idx, header in enumerate(headers, 1):
        cell = ws.cell(row=1, column=col_idx, value=header)
        cell.font = font
        cell.fill = fill
        cell.border = border
        cell.alignment = Alignment(horizontal="center", wrap_text=True)


def _xlsx_set_widths(ws, widths: list) -> None:
    from openpyxl.utils import get_column_letter  # type: ignore[import-untyped]

    for col_idx, width in enumerate(widths, 1):
        ws.column_dimensions[get_column_letter(col_idx)].width = width


# ── program register exports ──────────────────────────────────────────────

_PROGRAM_HEADERS = [
    "Program ID",
    "Title",
    "Category",
    "Status",
    "Priority",
    "Program manager",
    "Directorate",
    "Start date",
    "End date",
    "Approved budget",
    "Utilized budget",
    "Currency",
]
_PROGRAM_WIDTHS = [22, 32, 22, 14, 10, 22, 22, 14, 14, 18, 18, 12]


def _program_data_row(program) -> list:
    manager = program.program_manager
    return [
        program.reference_number,
        program.title,
        program.category.name if program.category else "",
        program.get_status_display(),
        program.get_priority_display(),
        manager.full_name if manager else "",
        program.responsible_directorate.name if program.responsible_directorate else "",
        program.start_date,
        program.end_date,
        program.budget_approved,
        program.budget_utilized,
        program.currency,
    ]


def program_register_xlsx_response(user) -> HttpResponse:
    from openpyxl import Workbook  # type: ignore[import-untyped]
    from openpyxl.styles import Border, Side

    if not (
        user_has_permission(user, PROGRAMMES_EXPORT)
        or user_has_permission(user, PROGRAMMES_MANAGE)
    ):
        raise PermissionDenied
    border = Border(
        left=Side(style="thin"),
        right=Side(style="thin"),
        top=Side(style="thin"),
        bottom=Side(style="thin"),
    )
    wb = Workbook()
    ws = wb.active
    ws.title = "Program Register"
    _xlsx_styled_header(ws, _PROGRAM_HEADERS)
    for row_idx, program in enumerate(
        visible_programs(user).order_by("reference_number").iterator(chunk_size=500), 2
    ):
        for col_idx, value in enumerate(_program_data_row(program), 1):
            cell = ws.cell(row=row_idx, column=col_idx, value=value)
            cell.border = border
    _xlsx_set_widths(ws, _PROGRAM_WIDTHS)
    buf = BytesIO()
    wb.save(buf)
    buf.seek(0)
    resp = HttpResponse(
        buf.getvalue(),
        content_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
    )
    resp["Content-Disposition"] = 'attachment; filename="program_register.xlsx"'
    resp["Cache-Control"] = "private, no-store"
    logger.info(
        "program_register_xlsx_exported",
        extra={"program_event": {"action": "export", "format": "xlsx"}},
    )
    return resp


def program_register_docx_response(user) -> HttpResponse:
    from docx import Document as DocxDocument

    if not (
        user_has_permission(user, PROGRAMMES_EXPORT)
        or user_has_permission(user, PROGRAMMES_MANAGE)
    ):
        raise PermissionDenied
    doc = DocxDocument()
    doc.add_heading("Program Register", 0)
    table = doc.add_table(rows=1, cols=len(_PROGRAM_HEADERS))
    table.style = "Table Grid"
    for cell, label in zip(table.rows[0].cells, _PROGRAM_HEADERS, strict=False):
        cell.text = label
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True
    for program in (
        visible_programs(user).order_by("reference_number").iterator(chunk_size=500)
    ):
        row_cells = table.add_row().cells
        for cell, value in zip(row_cells, _program_data_row(program), strict=False):
            cell.text = str(value) if value is not None else ""
    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    resp = HttpResponse(
        buf.getvalue(),
        content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )
    resp["Content-Disposition"] = 'attachment; filename="program_register.docx"'
    resp["Cache-Control"] = "private, no-store"
    logger.info(
        "program_export_docx_exported",
        extra={"program_event": {"actor": str(user.pk), "format": "docx"}},
    )
    return resp


def program_register_pdf_response(user) -> HttpResponse:
    from reportlab.lib.colors import HexColor  # type: ignore[import-untyped]
    from reportlab.lib.pagesizes import A4  # type: ignore[import-untyped]
    from reportlab.lib.units import mm  # type: ignore[import-untyped]
    from reportlab.platypus import (  # type: ignore[import-untyped]
        Image,
        SimpleDocTemplate,
        Spacer,
        Table,
        TableStyle,
    )

    if not (
        user_has_permission(user, PROGRAMMES_EXPORT)
        or user_has_permission(user, PROGRAMMES_MANAGE)
    ):
        raise PermissionDenied
    buf = BytesIO()
    doc = SimpleDocTemplate(
        buf,
        pagesize=A4,
        rightMargin=15 * mm,
        leftMargin=15 * mm,
        topMargin=15 * mm,
        bottomMargin=15 * mm,
    )
    story = []
    if LOGO_PATH.is_file():
        story.append(Image(str(LOGO_PATH), width=36 * mm, height=24 * mm))
        story.append(Spacer(1, 6 * mm))
    story.append(
        Table(
            [["Program Register"]],
            style=TableStyle(
                [
                    ("FONTSIZE", (0, 0), (-1, -1), 16),
                    ("ALIGN", (0, 0), (-1, -1), "CENTER"),
                ]
            ),
        )
    )
    story.append(Spacer(1, 4 * mm))
    data = [_PROGRAM_HEADERS]
    for program in (
        visible_programs(user).order_by("reference_number").iterator(chunk_size=500)
    ):
        data.append(
            [str(v) if v is not None else "" for v in _program_data_row(program)]
        )
    t = Table(data, repeatRows=1)
    t.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, 0), HexColor("#003366")),
                ("TEXTCOLOR", (0, 0), (-1, 0), HexColor("#ffffff")),
                ("GRID", (0, 0), (-1, -1), 0.5, HexColor("#cccccc")),
                ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                ("FONTSIZE", (0, 0), (-1, -1), 7),
            ]
        )
    )
    story.append(t)
    doc.build(story)
    buf.seek(0)
    resp = HttpResponse(buf.getvalue(), content_type="application/pdf")
    resp["Content-Disposition"] = 'attachment; filename="program_register.pdf"'
    resp["Cache-Control"] = "private, no-store"
    logger.info(
        "program_export_pdf_exported",
        extra={"program_event": {"actor": str(user.pk), "format": "pdf"}},
    )
    return resp


_PROJECT_HEADERS = [
    "Project ID",
    "Program ID",
    "Title",
    "Category",
    "Status",
    "Project Manager",
    "Start date",
    "End date",
    "Approved budget",
    "Utilized budget",
    "Currency",
]
_PROJECT_WR = [22, 22, 32, 22, 14, 22, 14, 14, 18, 18, 12]


def _project_data_row(project) -> list:
    manager = project.project_manager
    return [
        project.reference_number,
        project.program.reference_number,
        project.title,
        project.category.name if project.category else "",
        project.get_status_display(),
        manager.full_name if manager else "",
        project.start_date,
        project.end_date,
        project.budget_approved,
        project.budget_utilized,
        project.currency,
    ]


def project_register_xlsx_response(user) -> HttpResponse:
    from openpyxl import Workbook
    from openpyxl.styles import Border, Side

    if not (
        user_has_permission(user, PROJECTS_EXPORT)
        or user_has_permission(user, PROJECTS_MANAGE)
    ):
        raise PermissionDenied
    border = Border(
        left=Side(style="thin"),
        right=Side(style="thin"),
        top=Side(style="thin"),
        bottom=Side(style="thin"),
    )
    wb = Workbook()
    ws = wb.active
    ws.title = "Project Register"
    _xlsx_styled_header(ws, _PROJECT_HEADERS)
    for row_idx, project in enumerate(
        visible_projects(user).order_by("reference_number").iterator(chunk_size=500), 2
    ):
        for col_idx, value in enumerate(_project_data_row(project), 1):
            cell = ws.cell(row=row_idx, column=col_idx, value=value)
            cell.border = border
    _xlsx_set_widths(ws, _PROJECT_WR)
    buf = BytesIO()
    wb.save(buf)
    buf.seek(0)
    resp = HttpResponse(
        buf.getvalue(),
        content_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
    )
    resp["Content-Disposition"] = 'attachment; filename="project_register.xlsx"'
    resp["Cache-Control"] = "private, no-store"
    logger.info(
        "project_export_xlsx",
        extra={"program_event": {"actor": str(user.pk), "format": "xlsx"}},
    )
    return resp


def project_register_docx_response(user) -> HttpResponse:
    from docx import Document as DocxDocument

    if not (
        user_has_permission(user, PROJECTS_EXPORT)
        or user_has_permission(user, PROJECTS_MANAGE)
    ):
        raise PermissionDenied
    doc = DocxDocument()
    doc.add_heading("Project Register", 0)
    table = doc.add_table(rows=1, cols=len(_PROJECT_HEADERS))
    table.style = "Table Grid"
    for cell, label in zip(table.rows[0].cells, _PROJECT_HEADERS, strict=False):
        cell.text = label
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True
    for project in (
        visible_projects(user).order_by("reference_number").iterator(chunk_size=500)
    ):
        row_cells = table.add_row().cells
        for cell, value in zip(row_cells, _project_data_row(project), strict=False):
            cell.text = str(value) if value is not None else ""
    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    resp = HttpResponse(
        buf.getvalue(),
        content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )
    resp["Content-Disposition"] = 'attachment; filename="project_register.docx"'
    resp["Cache-Control"] = "private, no-store"
    return resp


def project_register_pdf_response(user) -> HttpResponse:
    if not (
        user_has_permission(user, PROJECTS_EXPORT)
        or user_has_permission(user, PROJECTS_MANAGE)
    ):
        raise PermissionDenied
    from reportlab.lib.colors import HexColor
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.units import mm
    from reportlab.platypus import Image, SimpleDocTemplate, Spacer, Table, TableStyle

    buf = BytesIO()
    doc = SimpleDocTemplate(
        buf,
        pagesize=A4,
        rightMargin=15 * mm,
        leftMargin=15 * mm,
        topMargin=15 * mm,
        bottomMargin=15 * mm,
    )
    story = []
    if LOGO_PATH.is_file():
        story.append(Image(str(LOGO_PATH), width=36 * mm, height=24 * mm))
        story.append(Spacer(1, 6 * mm))
    story.append(
        Table(
            [["Project Register"]],
            style=TableStyle(
                [
                    ("FONTSIZE", (0, 0), (-1, -1), 16),
                    ("ALIGN", (0, 0), (-1, -1), "CENTER"),
                ]
            ),
        )
    )
    story.append(Spacer(1, 4 * mm))
    data = [_PROJECT_HEADERS]
    for project in (
        visible_projects(user).order_by("reference_number").iterator(chunk_size=500)
    ):
        data.append(
            [str(v) if v is not None else "" for v in _project_data_row(project)]
        )
    t = Table(data, repeatRows=1)
    t.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, 0), HexColor("#003366")),
                ("TEXTCOLOR", (0, 0), (-1, 0), HexColor("#ffffff")),
                ("GRID", (0, 0), (-1, -1), 0.5, HexColor("#cccccc")),
                ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                ("FONTSIZE", (0, 0), (-1, -1), 7),
            ]
        )
    )
    story.append(t)
    doc.build(story)
    buf.seek(0)
    resp = HttpResponse(buf.getvalue(), content_type="application/pdf")
    resp["Content-Disposition"] = 'attachment; filename="project_register.pdf"'
    resp["Cache-Control"] = "private, no-store"
    return resp


# ── lessons learned exports ───────────────────────────────────────────────

_LESSON_HEADERS = [
    "Lesson ID",
    "Program",
    "Title",
    "Category",
    "Summary",
    "Recorded at",
]
_LESSON_WIDTHS = [21, 21, 28, 18, 40, 14]


def _lesson_data_row(lesson) -> list:
    return [
        lesson.reference_number,
        lesson.program.title,
        lesson.title,
        lesson.get_category_display(),
        lesson.summary[:200],
        lesson.recorded_at,
    ]


def lessons_learned_xlsx_response(user) -> HttpResponse:
    from openpyxl import Workbook
    from openpyxl.styles import Border, Side

    if not (
        user_has_permission(user, PROGRAMMES_EXPORT)
        or user_has_permission(user, PROGRAMMES_MANAGE)
    ):
        raise PermissionDenied
    from .models import LessonsLearned

    border = Border(
        left=Side(style="thin"),
        right=Side(style="thin"),
        top=Side(style="thin"),
        bottom=Side(style="thin"),
    )
    wb = Workbook()
    ws = wb.active
    ws.title = "Lessons Learned"
    headers = _LESSON_HEADERS
    _xlsx_styled_header(ws, headers)
    rows = LessonsLearned.objects.select_related("program").order_by("-recorded_at")
    for row_idx, lesson in enumerate(rows.iterator(chunk_size=500), 2):
        for col_idx, value in enumerate(_lesson_data_row(lesson), 1):
            cell = ws.cell(row=row_idx, column=col_idx, value=value)
            cell.border = border
    _xlsx_set_widths(ws, _LESSON_WIDTHS)
    buf = BytesIO()
    wb.save(buf)
    buf.seek(0)
    resp = HttpResponse(
        buf.getvalue(),
        content_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
    )
    resp["Content-Disposition"] = "attachment;filename=lessons_learned.xlsx"
    resp["Cache-Control"] = "private, no-store"
    return resp


# ── closure report docx helper ────────────────────────────────────────────


def program_closure_docx_response(user, program_id: str) -> HttpResponse:
    from docx import Document as DocxDocument

    if not (
        user_has_permission(user, PROGRAMMES_EXPORT)
        or user_has_permission(user, PROGRAMMES_MANAGE)
    ):
        raise PermissionDenied
    from .models import Program

    program = Program.objects.get(pk=program_id)
    doc = DocxDocument()
    doc.add_heading("Program Closure Report", 0)
    doc.add_heading(program.title, 1)
    doc.add_paragraph(f"Program ID: {program.reference_number}")
    doc.add_paragraph(f"Status: {program.get_status_display()}")
    program_manager = (
        program.program_manager.full_name if program.program_manager else "N/A"
    )
    doc.add_paragraph(f"Program Manager: {program_manager}")
    doc.add_paragraph(f"Period: {program.start_date} to {program.end_date}")
    doc.add_paragraph(
        f"Budget: {program.budget_utilized}/{program.budget_approved} "
        f"{program.currency}"
    )
    doc.add_paragraph("")
    doc.add_heading("Beneficiary Reach", 2)
    doc.add_paragraph(f"Target: {program.target_beneficiaries}")
    doc.add_paragraph(f"Count: {program.target_beneficiary_count or 'Not set'}")
    doc.add_paragraph("")
    doc.add_heading("Lessons Learned", 2)
    for lesson in program.lessons_learned.all()[:20]:
        doc.add_paragraph(
            f"[{lesson.get_category_display()}] {lesson.title}: {lesson.summary[:200]}"
        )
    doc.add_paragraph("")
    doc.add_heading("Evaluations", 2)
    for evaluation in program.evaluations.select_related("conducted_by")[:10]:
        doc.add_paragraph(
            f"({evaluation.get_evaluation_type_display()}) {evaluation.title} — "
            f"{evaluation.evaluation_date}"
        )
        if evaluation.findings:
            doc.add_paragraph(evaluation.findings[:500])
    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    resp = HttpResponse(
        buf.getvalue(),
        content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )
    resp["Content-Disposition"] = (
        f"""attachment; filename="program_{program.reference_number}_closure.docx" """
    )
    resp["Cache-Control"] = "private, no-store"
    return resp
