"""Multi-format export adapters for Calendar & Meetings.

CSV/XLSX/DOCX/PDF adapters mirror ``apps.registers.exports``.  Every export
is permission checked and records an immutable activity timeline entry.
"""

from __future__ import annotations

import csv
import json
import logging
from io import BytesIO, StringIO
from pathlib import Path

from django.conf import settings
from django.core.exceptions import PermissionDenied
from django.http import HttpResponse

from apps.rbac.authorization import user_has_permission

from .models import MeetingActivityRecord
from .permissions import MEETING_EXPORT, MEETING_MANAGE
from .selectors import visible_calendars, visible_events, visible_meetings

logger = logging.getLogger(__name__)

LOGO_PATH = Path(settings.BASE_DIR) / "static" / "images" / "app_logo.png"


def formula_safe_csv_value(value) -> str:
    text = str(value or "")
    if text.startswith(("=", "+", "-", "@", "\t", "\r")):
        return f"'{text}"
    return text


def _require_export_permission(user) -> None:
    if not (
        user_has_permission(user, MEETING_EXPORT)
        or user_has_permission(user, MEETING_MANAGE)
    ):
        raise PermissionDenied("Meeting export permission is required.")


def _csv_response(user, headers, rows, filename: str, event: str) -> HttpResponse:
    output = StringIO(newline="")
    writer = csv.writer(output)
    writer.writerow(headers)
    for row in rows:
        writer.writerow([formula_safe_csv_value(value) for value in row])
    response = HttpResponse(output.getvalue(), content_type="text/csv; charset=utf-8")
    response["Content-Disposition"] = f'attachment; filename="{filename}"'
    response["Cache-Control"] = "private, no-store"
    response["Pragma"] = "no-cache"
    response["X-Content-Type-Options"] = "nosniff"
    logger.info(
        event,
        extra={
            "meeting_event": {
                "action": event,
                "format": "csv",
                "actor_id": str(user.pk),
                "row_count": len(rows),
                "scope": "visible_meetings",
            }
        },
    )
    return response


def _xlsx_styled_header(sheet, headers) -> None:
    from openpyxl.styles import Font, PatternFill  # type: ignore[import-untyped]

    fill = PatternFill(start_color="003366", end_color="003366", fill_type="solid")
    for col_idx, header in enumerate(headers, 1):
        cell = sheet.cell(row=1, column=col_idx, value=header)
        cell.font = Font(bold=True, color="FFFFFF")
        cell.fill = fill


def _xlsx_set_widths(sheet, widths: list) -> None:
    from openpyxl.utils import get_column_letter  # type: ignore[import-untyped]

    for col_idx, width in enumerate(widths, 1):
        sheet.column_dimensions[get_column_letter(col_idx)].width = width


def _build_xlsx(title: str, rows: list[list], headers: list, widths: list) -> BytesIO:
    from openpyxl import Workbook  # type: ignore[import-untyped]

    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "Calendar & Meetings"
    _xlsx_styled_header(sheet, headers)
    for row in rows:
        sheet.append(row)
    _xlsx_set_widths(sheet, widths)
    buffer = BytesIO()
    workbook.save(buffer)
    buffer.seek(0)
    return buffer


def _build_docx(title: str, rows: list[list], headers: list) -> BytesIO:
    from docx import Document

    document = Document()
    document.add_heading(title, level=0)
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Light Grid Accent 1"
    for col_idx, header in enumerate(headers):
        table.rows[0].cells[col_idx].text = str(header)
    for row in rows:
        cells = table.add_row().cells
        for col_idx, value in enumerate(row):
            cells[col_idx].text = str(value)
    buffer = BytesIO()
    document.save(buffer)
    buffer.seek(0)
    return buffer


def _build_pdf(title: str, rows: list[list], headers: list) -> BytesIO:
    from reportlab.lib import colors  # type: ignore[import-untyped]
    from reportlab.lib.pagesizes import A4, landscape  # type: ignore[import-untyped]
    from reportlab.lib.styles import getSampleStyleSheet  # type: ignore[import-untyped]
    from reportlab.platypus import (  # type: ignore[import-untyped]
        SimpleDocTemplate,
        Table,
        TableStyle,
    )

    buffer = BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=landscape(A4), title=title)
    styles = getSampleStyleSheet()
    elements = [_reportlab_heading(title, styles)]
    data = [headers] + [[str(value) for value in row] for row in rows]
    table = Table(data, repeatRows=1)
    table.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#003366")),
                ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                ("FONTSIZE", (0, 0), (-1, 0), 8),
                ("FONTSIZE", (0, 1), (-1, -1), 7),
                ("GRID", (0, 0), (-1, -1), 0.25, colors.grey),
                (
                    "ROWBACKGROUNDS",
                    (0, 1),
                    (-1, -1),
                    [colors.white, colors.HexColor("#EEF2F7")],
                ),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
            ]
        )
    )
    elements.append(table)
    doc.build(elements)
    buffer.seek(0)
    return buffer


def _reportlab_heading(title: str, styles) -> object:
    from reportlab.lib import colors
    from reportlab.lib.styles import ParagraphStyle
    from reportlab.platypus import Paragraph

    style = ParagraphStyle(
        "MeetingsTitle",
        parent=styles["Title"],
        fontName="Helvetica-Bold",
        fontSize=14,
        leading=18,
        textColor=colors.HexColor("#003366"),
    )
    return Paragraph(title, style)


MEETING_HEADERS = [
    "Reference",
    "Title",
    "Meeting Type",
    "Status",
    "Start",
    "End",
    "Mode",
    "Venue",
    "Organizer",
    "Program",
    "Project",
    "Organization Unit",
    "Confidentiality",
    "Quorum Type",
    "Quorum Value",
    "Agenda Status",
    "Minutes Status",
    "Created At",
    "Updated At",
]


def _meeting_row(meeting) -> list:
    start_at = meeting.start_at
    if start_at is not None and getattr(start_at, "tzinfo", None) is not None:
        start_at = start_at.replace(tzinfo=None)
    end_at = meeting.end_at
    if end_at is not None and getattr(end_at, "tzinfo", None) is not None:
        end_at = end_at.replace(tzinfo=None)
    created_at = meeting.created_at
    if created_at is not None and getattr(created_at, "tzinfo", None) is not None:
        created_at = created_at.replace(tzinfo=None)
    updated_at = meeting.updated_at
    if updated_at is not None and getattr(updated_at, "tzinfo", None) is not None:
        updated_at = updated_at.replace(tzinfo=None)
    return [
        meeting.reference,
        meeting.title,
        meeting.get_meeting_type_display(),
        meeting.get_status_display(),
        start_at,
        end_at,
        meeting.get_mode_display(),
        meeting.venue.name if meeting.venue else "",
        meeting.organizer.full_name if meeting.organizer else "",
        meeting.program.title if meeting.program else "",
        meeting.project.title if meeting.project else "",
        meeting.organization_unit.name if meeting.organization_unit else "",
        meeting.get_confidentiality_level_display(),
        meeting.get_quorum_type_display(),
        meeting.quorum_value,
        meeting.get_agenda_status_display(),
        meeting.get_minutes_status_display(),
        created_at,
        updated_at,
    ]


def _record_export_activity(user, meeting=None, fmt: str = "csv") -> None:
    if meeting is not None:
        MeetingActivityRecord.objects.create(
            meeting=meeting,
            action="EXPORTED",
            actor=user,
            details=f"Meeting data exported as {fmt.upper()}.",
        )


def _exported_meetings(user, meeting=None):
    queryset = visible_meetings(user)
    if meeting is not None:
        queryset = queryset.filter(pk=meeting.pk)
    return queryset.select_related(
        "venue", "organizer", "program", "project", "organization_unit"
    )


def meetings_export_response(user, meeting=None, fmt: str = "csv") -> HttpResponse:
    """Export visible meetings (or a single meeting) in the requested format."""
    _require_export_permission(user)
    entries = list(_exported_meetings(user, meeting=meeting))
    title = "Calendar & Meetings Export"
    rows = [_meeting_row(entry) for entry in entries]
    widths = [
        18,
        30,
        20,
        20,
        18,
        18,
        16,
        24,
        22,
        22,
        22,
        22,
        18,
        18,
        14,
        18,
        20,
        22,
        22,
    ]

    if fmt == "csv":
        _record_export_activity(user, meeting=meeting, fmt="csv")
        return _csv_response(
            user,
            MEETING_HEADERS,
            rows,
            "meetings_export.csv",
            "meetings_exported",
        )
    if fmt == "json":
        _record_export_activity(user, meeting=meeting, fmt="json")
        safe_rows = [[formula_safe_csv_value(v) for v in row] for row in rows]
        payload = [dict(zip(MEETING_HEADERS, row, strict=False)) for row in safe_rows]
        response = HttpResponse(
            json.dumps(payload, indent=2, default=str),
            content_type="application/json; charset=utf-8",
        )
        response["Content-Disposition"] = 'attachment; filename="meetings_export.json"'
        response["Cache-Control"] = "private, no-store"
        response["X-Content-Type-Options"] = "nosniff"
        return response

    if fmt == "xlsx":
        _record_export_activity(user, meeting=meeting, fmt="xlsx")
        buffer = _build_xlsx(title, rows, MEETING_HEADERS, widths)
        content_type = (
            "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
        )
        filename = "meetings_export.xlsx"
    elif fmt == "docx":
        _record_export_activity(user, meeting=meeting, fmt="docx")
        buffer = _build_docx(title, rows, MEETING_HEADERS)
        content_type = (
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
        filename = "meetings_export.docx"
    elif fmt == "pdf":
        _record_export_activity(user, meeting=meeting, fmt="pdf")
        buffer = _build_pdf(title, rows, MEETING_HEADERS)
        content_type = "application/pdf"
        filename = "meetings_export.pdf"
    else:
        raise PermissionDenied(f"Unsupported export format: {fmt}")

    response = HttpResponse(buffer.getvalue(), content_type=content_type)
    response["Content-Disposition"] = f'attachment; filename="{filename}"'
    response["Cache-Control"] = "private, no-store"
    response["Pragma"] = "no-cache"
    response["X-Content-Type-Options"] = "nosniff"
    return response


def calendars_export_response(user, fmt: str = "csv") -> HttpResponse:
    """Export visible calendars in the requested format."""
    _require_export_permission(user)
    entries = list(visible_calendars(user))
    headers = ["Name", "Description", "Type", "Confidentiality", "Owner", "Created At"]
    rows = []
    for calendar in entries:
        created_at = calendar.created_at
        if created_at is not None and getattr(created_at, "tzinfo", None) is not None:
            created_at = created_at.replace(tzinfo=None)
        rows.append(
            [
                calendar.name,
                calendar.description,
                calendar.get_calendar_type_display(),
                calendar.get_confidentiality_level_display(),
                calendar.owner.full_name if calendar.owner else "",
                created_at,
            ]
        )
    if fmt == "csv":
        return _csv_response(
            user,
            headers,
            rows,
            "calendars_export.csv",
            "calendars_exported",
        )
    raise PermissionDenied(f"Unsupported export format: {fmt}")


def events_export_response(user, fmt: str = "csv") -> HttpResponse:
    """Export visible events in the requested format."""
    _require_export_permission(user)
    entries = list(visible_events(user))
    headers = [
        "Reference",
        "Title",
        "Event Type",
        "Start",
        "End",
        "Calendar",
        "Venue",
        "Host",
        "Organizer",
        "Program",
        "Project",
        "Status",
        "Confidentiality",
    ]
    rows = []
    for event in entries:
        start_at = event.start_at
        if start_at is not None and getattr(start_at, "tzinfo", None) is not None:
            start_at = start_at.replace(tzinfo=None)
        end_at = event.end_at
        if end_at is not None and getattr(end_at, "tzinfo", None) is not None:
            end_at = end_at.replace(tzinfo=None)
        rows.append(
            [
                event.reference,
                event.title,
                event.get_event_type_display(),
                start_at,
                end_at,
                event.calendar.name if event.calendar else "",
                event.venue.name if event.venue else "",
                event.host.full_name if event.host else "",
                event.organizer.full_name if event.organizer else "",
                event.program.title if event.program else "",
                event.project.title if event.project else "",
                event.get_status_display(),
                event.get_confidentiality_level_display(),
            ]
        )
    if fmt == "csv":
        return _csv_response(
            user,
            headers,
            rows,
            "events_export.csv",
            "events_exported",
        )
    raise PermissionDenied(f"Unsupported export format: {fmt}")
