"""Multi-format (XLSX, DOCX, PDF) report export adapters for beneficiaries."""

from __future__ import annotations

import logging
from io import BytesIO
from pathlib import Path

from django.conf import settings
from django.core.exceptions import PermissionDenied
from django.http import HttpResponse

from apps.rbac.authorization import user_has_permission

from .permissions import BENEFICIARIES_EXPORT, BENEFICIARIES_MANAGE
from .selectors import visible_beneficiaries

logger = logging.getLogger(__name__)

LOGO_PATH = Path(settings.BASE_DIR) / "static" / "images" / "app_logo.png"


def _xlsx_styled_header(ws, headers: list) -> None:
    from openpyxl.styles import (  # type: ignore[import-untyped]
        Alignment,
        Border,
        Font,
        PatternFill,
        Side,
    )

    fill = PatternFill(start_color="003366", end_color="003366", fill_type="solid")
    font = Font(color="FFFFFF", bold=True, size=11)
    border = Border(
        left=Side(style="thin"),
        right=Side(style="thin"),
        top=Side(style="thin"),
        bottom=Side(style="thin"),
    )
    for col_idx, header in enumerate(headers, 1):
        cell = ws.cell(row=1, column=col_idx, value=header)
        cell.font = font
        cell.fill = fill
        cell.border = border
        cell.alignment = Alignment(horizontal="center", wrap_text=True)


def _xlsx_set_widths(ws, widths: list) -> None:
    from openpyxl.utils import get_column_letter  # type: ignore[import-untyped]

    for col_idx, width in enumerate(widths, 1):
        ws.column_dimensions[get_column_letter(col_idx)].width = width


_BENEFICIARY_HEADERS = [
    "Beneficiary ID",
    "Full name",
    "Category",
    "Classification",
    "Status",
    "Confidentiality",
    "Gender",
    "Date of birth",
    "Is minor",
    "Phone",
    "Province or region",
    "District",
    "Community",
    "Household",
    "Responsible officer",
    "Case manager",
    "Registration date",
    "Consent status",
]
_BENEFICIARY_WIDTHS = [
    22,
    30,
    22,
    18,
    14,
    16,
    14,
    14,
    10,
    18,
    18,
    18,
    18,
    22,
    22,
    22,
    14,
    16,
]


def _beneficiary_data_row(beneficiary) -> list:
    responsible_officer = beneficiary.primary_responsible_officer
    case_manager = beneficiary.case_manager
    household = beneficiary.household
    return [
        beneficiary.reference_number,
        beneficiary.full_name,
        beneficiary.category.name if beneficiary.category else "",
        beneficiary.classification.name if beneficiary.classification else "",
        beneficiary.get_status_display(),
        beneficiary.get_confidentiality_display(),
        beneficiary.gender.name if beneficiary.gender else "",
        beneficiary.date_of_birth,
        "Yes" if beneficiary.is_minor else "No",
        beneficiary.phone_primary,
        beneficiary.province_or_region,
        beneficiary.district,
        beneficiary.community,
        household.reference_number if household else "",
        responsible_officer.full_name if responsible_officer else "",
        case_manager.full_name if case_manager else "",
        beneficiary.registration_date,
        beneficiary.get_consent_status_display(),
    ]


def _register_guard(user) -> None:
    if not (
        user_has_permission(user, BENEFICIARIES_EXPORT)
        or user_has_permission(user, BENEFICIARIES_MANAGE)
    ):
        raise PermissionDenied


def beneficiary_register_xlsx_response(user) -> HttpResponse:
    from openpyxl import Workbook  # type: ignore[import-untyped]
    from openpyxl.styles import Border, Side

    _register_guard(user)
    border = Border(
        left=Side(style="thin"),
        right=Side(style="thin"),
        top=Side(style="thin"),
        bottom=Side(style="thin"),
    )
    wb = Workbook()
    ws = wb.active
    ws.title = "Beneficiary Register"
    _xlsx_styled_header(ws, _BENEFICIARY_HEADERS)
    for row_idx, beneficiary in enumerate(
        visible_beneficiaries(user)
        .order_by("reference_number")
        .iterator(chunk_size=500),
        2,
    ):
        for col_idx, value in enumerate(_beneficiary_data_row(beneficiary), 1):
            cell = ws.cell(row=row_idx, column=col_idx, value=value)
            cell.border = border
    _xlsx_set_widths(ws, _BENEFICIARY_WIDTHS)
    buf = BytesIO()
    wb.save(buf)
    buf.seek(0)
    resp = HttpResponse(
        buf.getvalue(),
        content_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
    )
    resp["Content-Disposition"] = 'attachment; filename="beneficiary_register.xlsx"'
    resp["Cache-Control"] = "private, no-store"
    logger.info("beneficiary_register_xlsx_exported")
    return resp


def beneficiary_register_docx_response(user) -> HttpResponse:
    from docx import Document as DocxDocument

    _register_guard(user)
    doc = DocxDocument()
    doc.add_heading("Beneficiary Register", 0)
    table = doc.add_table(rows=1, cols=len(_BENEFICIARY_HEADERS))
    table.style = "Table Grid"
    for cell, label in zip(table.rows[0].cells, _BENEFICIARY_HEADERS, strict=False):
        cell.text = label
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True
    for beneficiary in (
        visible_beneficiaries(user)
        .order_by("reference_number")
        .iterator(chunk_size=500)
    ):
        row_cells = table.add_row().cells
        for cell, value in zip(
            row_cells, _beneficiary_data_row(beneficiary), strict=False
        ):
            cell.text = str(value) if value is not None else ""
    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    resp = HttpResponse(
        buf.getvalue(),
        content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )
    resp["Content-Disposition"] = 'attachment; filename="beneficiary_register.docx"'
    resp["Cache-Control"] = "private, no-store"
    return resp


def beneficiary_register_pdf_response(user) -> HttpResponse:
    from reportlab.lib.colors import HexColor  # type: ignore[import-untyped]
    from reportlab.lib.pagesizes import A4  # type: ignore[import-untyped]
    from reportlab.lib.units import mm  # type: ignore[import-untyped]
    from reportlab.platypus import (  # type: ignore[import-untyped]
        Image,
        SimpleDocTemplate,
        Spacer,
        Table,
        TableStyle,
    )

    _register_guard(user)
    buf = BytesIO()
    doc = SimpleDocTemplate(
        buf,
        pagesize=A4,
        rightMargin=15 * mm,
        leftMargin=15 * mm,
        topMargin=15 * mm,
        bottomMargin=15 * mm,
    )
    story = []
    if LOGO_PATH.is_file():
        story.append(Image(str(LOGO_PATH), width=36 * mm, height=24 * mm))
        story.append(Spacer(1, 6 * mm))
    story.append(
        Table(
            [["Beneficiary Register"]],
            style=TableStyle(
                [
                    ("FONTSIZE", (0, 0), (-1, -1), 16),
                    ("ALIGN", (0, 0), (-1, -1), "CENTER"),
                ]
            ),
        )
    )
    story.append(Spacer(1, 4 * mm))
    data = [_BENEFICIARY_HEADERS]
    for beneficiary in (
        visible_beneficiaries(user)
        .order_by("reference_number")
        .iterator(chunk_size=500)
    ):
        data.append(
            [
                str(v) if v is not None else ""
                for v in _beneficiary_data_row(beneficiary)
            ]
        )
    t = Table(data, repeatRows=1)
    t.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, 0), HexColor("#003366")),
                ("TEXTCOLOR", (0, 0), (-1, 0), HexColor("#ffffff")),
                ("GRID", (0, 0), (-1, -1), 0.5, HexColor("#cccccc")),
                ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                ("FONTSIZE", (0, 0), (-1, -1), 7),
            ]
        )
    )
    story.append(t)
    doc.build(story)
    buf.seek(0)
    resp = HttpResponse(buf.getvalue(), content_type="application/pdf")
    resp["Content-Disposition"] = 'attachment; filename="beneficiary_register.pdf"'
    resp["Cache-Control"] = "private, no-store"
    return resp


def beneficiary_profile_docx_response(user, beneficiary) -> HttpResponse:
    from docx import Document as DocxDocument

    if not (
        user_has_permission(user, BENEFICIARIES_EXPORT)
        or user_has_permission(user, BENEFICIARIES_MANAGE)
    ):
        raise PermissionDenied
    from .selectors import user_can_access_beneficiary

    if not user_can_access_beneficiary(user, beneficiary, include_archived=True):
        raise PermissionDenied

    responsible_officer = beneficiary.primary_responsible_officer
    case_manager = beneficiary.case_manager
    household = beneficiary.household
    rows = [
        ("Beneficiary ID", beneficiary.reference_number),
        ("Full name", beneficiary.full_name),
        ("Status", beneficiary.get_status_display()),
        ("Category", beneficiary.category.name if beneficiary.category else "N/A"),
        (
            "Classification",
            beneficiary.classification.name if beneficiary.classification else "N/A",
        ),
        ("Gender", beneficiary.gender.name if beneficiary.gender else "N/A"),
        ("Date of birth", beneficiary.date_of_birth or "N/A"),
        ("Is minor", "Yes" if beneficiary.is_minor else "No"),
        ("Phone", beneficiary.phone_primary),
        ("Email", beneficiary.email),
        ("Province or region", beneficiary.province_or_region),
        ("District", beneficiary.district),
        ("Community", beneficiary.community),
        (
            "Household",
            household.reference_number if household else "N/A",
        ),
        (
            "Responsible officer",
            responsible_officer.full_name if responsible_officer else "N/A",
        ),
        ("Case manager", case_manager.full_name if case_manager else "N/A"),
        ("Registration date", beneficiary.registration_date),
        ("Consent status", beneficiary.get_consent_status_display()),
        ("Consent expiry", beneficiary.consent_expiry_date or "N/A"),
    ]
    doc = DocxDocument()
    doc.add_heading("Beneficiary Profile", 0)
    doc.add_heading(beneficiary.full_name, 1)
    for field, value in rows:
        doc.add_paragraph(f"{field}: {value}")
    doc.add_heading("Vulnerabilities", 2)
    for item in beneficiary.vulnerabilities.all():
        doc.add_paragraph(f"- {item.name}")
    doc.add_heading("Needs", 2)
    for item in beneficiary.needs.all():
        doc.add_paragraph(f"- {item.name}")
    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    resp = HttpResponse(
        buf.getvalue(),
        content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )
    resp["Content-Disposition"] = (
        'attachment; filename="beneficiary_'
        f'{beneficiary.reference_number}_profile.docx"'
    )
    resp["Cache-Control"] = "private, no-store"
    return resp
